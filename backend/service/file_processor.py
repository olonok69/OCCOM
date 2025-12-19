import fitz  # PyMuPDF
import math
import os
import io
import logging
from typing import Dict, List, Tuple
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from markitdown import MarkItDown
import tempfile
import re


logger = logging.getLogger(__name__)


class PDFProcessor:
    def __init__(self, blob_service=None):
        """
        Initialize PDF processor with optional blob storage service.

        Args:
            blob_service: BlobStorageService instance for storing images
        """
        self.blob_service = blob_service

    def is_important_image(self, pdf_document, xref, page_width, page_height) -> bool:
        """
        Determine if an image is likely to be important content vs decorative elements.

        Args:
            pdf_document: PyMuPDF document object
            xref: Image XREF number
            page_width: Width of the page in points
            page_height: Height of the page in points

        Returns:
            True if image is likely important, False if likely decorative
        """
        try:
            base_image = pdf_document.extract_image(xref)
            img_width = base_image.get("width", 0)
            img_height = base_image.get("height", 0)

            # Skip very small images (likely icons or decorative elements)
            if img_width < 50 or img_height < 50:
                return False, base_image

            # Skip very thin/long images that might be decorative lines or borders
            aspect_ratio = (
                max(img_width, img_height) / min(img_width, img_height)
                if min(img_width, img_height) > 0
                else 0
            )
            if aspect_ratio > 20:  # Very elongated images
                return False, base_image

            return True, base_image

        except Exception as e:
            logger.warning(f"Warning: Error checking image importance: {e}")
            # If we can't determine, assume it's important to be safe
            return True, base_image

    def _take_page_screenshot(self, page, dpi=150):
        """
        Take a screenshot of a PDF page.

        Args:
            page: PyMuPDF page object
            dpi: Resolution for the screenshot (default: 150)

        Returns:
            bytes: PNG image data of the page screenshot
        """
        try:
            # Create a transformation matrix for the desired DPI
            mat = fitz.Matrix(dpi / 72, dpi / 72)  # 72 is the default DPI

            # Render page to pixmap
            pix = page.get_pixmap(matrix=mat)

            # Convert to PNG bytes
            png_bytes = pix.tobytes("png")

            return png_bytes

        except Exception as e:
            print(f"Error taking page screenshot: {e}")
            raise

    def iter_pdf(
        self, pdf_bytes: bytes, report_dir: str
    ) -> Tuple[Dict[int, Dict[str, str]], List[Tuple[str, bytes]]]:
        """
        Process PDF bytes and extract content by page and images.

        Args:
            pdf_bytes: PDF file content as bytes
            report_dir: Directory name for organizing extracted images

        Returns:
            Tuple of (content_by_page, image_batches)
            - content_by_page: Dict mapping page numbers to page content
            - image_batches: List of (image_path, image_bytes) tuples
        """
        content_by_page = {}
        image_batches = []

        try:
            # Open PDF from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")

            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                page_index = page_num + 1  # 1-indexed for user-facing page numbers

                # Extract text from page
                text = page.get_text()

                # Extract images from page
                image_list = page.get_images(full=True)
                page_images = []

                # Get page dimensions for image importance checking
                page_rect = page.rect
                page_width = page_rect.width
                page_height = page_rect.height

                # Check if page has important images
                has_important_images = False
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]  # Image XREF number
                    is_important, _ = self.is_important_image(
                        pdf_document, xref, page_width, page_height
                    )
                    if is_important:
                        has_important_images = True
                        break

                # If page has important images, take a complete page screenshot
                if has_important_images:
                    try:
                        screenshot_name = (
                            f"{report_dir}/page_{page_index}_screenshot.png"
                        )
                        screenshot_bytes = self._take_page_screenshot(page)

                        if self.blob_service:
                            try:
                                self.blob_service.upload_blob(
                                    screenshot_name, screenshot_bytes
                                )
                                page_images.append(screenshot_name)
                                print(
                                    f"Complete page screenshot captured for page {page_index}: {screenshot_name}"
                                )
                            except Exception as e:
                                logger.warning(
                                    f"Warning: Failed to upload image to blob storage: {e}"
                                )
                                image_batches.append(
                                    (screenshot_name, screenshot_bytes)
                                )
                                page_images.append(screenshot_name)
                        else:
                            image_batches.append((screenshot_name, screenshot_bytes))
                            page_images.append(screenshot_name)
                            print(f"Complete page screenshot saved: {screenshot_name}")

                    except Exception as e:
                        logger.warning(
                            f"Warning: Failed to extract image {img_index} from page {page_index}: {e}"
                        )

                # Store page content
                content_by_page[page_index] = {
                    "text": text,
                    "images": page_images,
                    "charts": [],  # Could be enhanced to detect charts
                    "tables": [],  # Could be enhanced with table extraction
                }

            pdf_document.close()

        except Exception as e:
            raise ValueError(f"Error processing PDF: {e}")

        return content_by_page, image_batches

    def chunk_text(self, text: str, words_per_chunk: int = 250) -> List[str]:
        """
        Split text into chunks of approximately the specified number of words.

        Args:
            text: Input text to chunk
            words_per_chunk: Target number of words per chunk

        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []

        words = text.split()
        chunks = []

        for i in range(0, len(words), words_per_chunk):
            chunk = " ".join(words[i : i + words_per_chunk])
            chunks.append(chunk)

        return chunks

    def clean_text(self, text: str) -> str:
        """
        Clean and normalize text content.

        Args:
            text: Input text to clean

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Basic text cleaning
        text = text.strip()
        # Remove excessive whitespace
        text = " ".join(text.split())

        return text

    def chunk_pdf_file_pages(self, file, chunk_size, chunk_index):
        """Yield a chunk of a PDF file by pages."""
        if not file.lower().endswith(".pdf"):
            raise ValueError("File must be a PDF")
        file_size = os.stat(file).st_size
        total_chunks = math.ceil(file_size / chunk_size)
        if chunk_index >= total_chunks:
            raise ValueError("Chunk index out of range")
        start_byte = chunk_index * chunk_size
        end_byte = min(start_byte + chunk_size, file_size)

        def generate_chunk():
            with open(file, "rb") as f:
                f.seek(start_byte)
                remaining_bytes = end_byte - start_byte
                while remaining_bytes > 0:
                    read_size = min(4096, remaining_bytes)
                    data = f.read(read_size)
                    if not data:
                        break
                    remaining_bytes -= len(data)
                    yield data

        return generate_chunk()

    def get_pdf_pages(self, file, start_page, end_page):
        """Extract text from a range of pages in a PDF file."""
        if not file.lower().endswith(".pdf"):
            raise ValueError("File must be a PDF")
        if start_page < 0 or end_page < start_page:
            raise ValueError("Invalid page range")
        try:
            pdf_document = fitz.open(file)
            num_pages = len(pdf_document)

            if end_page >= num_pages:
                raise ValueError("End page exceeds number of pages in the document")
            if end_page is None:
                end_page = start_page

            # Extract text from the specified page range
            text_content = []
            for page_num in range(start_page, end_page + 1):
                page = pdf_document[page_num]
                text_content.append(page.get_text())

            pdf_document.close()
            return "\n\n".join(text_content)

        except Exception as e:
            raise ValueError(f"Error reading PDF file: {e}")


class WordProcessor:
    """
    Processor for Microsoft Word (.docx) documents.
    Extracts text, tables, images, and charts with better structure preservation than PDF.
    Supports chapter-based chunking for structured documents.
    """

    def __init__(self):
        """Initialize the Word processor."""
        pass

    def iter_word_document(
        self, docx_bytes: bytes, report_dir: str
    ) -> Tuple[Dict[int, Dict[str, str]], List[Tuple[str, bytes]]]:
        """
        Process Word document bytes and extract structured content.

        Args:
            docx_bytes: Word document content as bytes
            report_dir: Directory name for organizing extracted images

        Returns:
            Tuple of (content_by_section, image_batches)
            - content_by_section: Dict mapping section numbers to content
            - image_batches: List of (image_path, image_bytes) tuples
        """
        content_by_section = {}
        image_batches = []

        try:
            # Create a file-like object from bytes
            docx_file = io.BytesIO(docx_bytes)
            document = Document(docx_file)

            # Extract content section by section
            section_num = 1
            current_section = {"text": [], "images": [], "charts": [], "tables": []}

            # Iterate through document elements in order
            for element in document.element.body:
                if isinstance(element, CT_P):
                    # Paragraph
                    paragraph = Paragraph(element, document)
                    text = paragraph.text.strip()

                    if text:
                        current_section["text"].append(text)

                    # Check for images in paragraph
                    images = self._extract_images_from_paragraph(
                        paragraph,
                        report_dir,
                        section_num,
                        len(current_section["images"]),
                    )
                    for img_path, img_bytes in images:
                        current_section["images"].append(img_path)
                        image_batches.append((img_path, img_bytes))

            # Store the final section
            content_by_section[section_num] = {
                "text": "\n".join(current_section["text"]),
                "images": current_section["images"],
                "charts": current_section["charts"],
                "tables": current_section["tables"],
            }

        except Exception as e:
            raise ValueError(f"Error processing Word document: {e}")

        return content_by_section, image_batches

    def _extract_images_from_paragraph(
        self, paragraph: Paragraph, report_dir: str, section_num: int, image_index: int
    ) -> List[Tuple[str, bytes]]:
        """
        Extract images from a paragraph.

        Args:
            paragraph: Word paragraph object
            report_dir: Directory name for organizing images
            section_num: Current section number
            image_index: Starting index for image numbering

        Returns:
            List of (image_path, image_bytes) tuples
        """
        images = []

        try:
            # Check for inline shapes (images)
            if hasattr(paragraph, "_element"):
                for run in paragraph.runs:
                    if hasattr(run, "_element"):
                        # Look for drawing elements (images)
                        for drawing in run._element.findall(
                            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                        ):
                            # Extract image blip
                            blips = drawing.findall(
                                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                            )
                            for blip in blips:
                                embed_id = blip.get(
                                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                )
                                if embed_id:
                                    # Get the image part
                                    image_part = paragraph.part.related_parts.get(
                                        embed_id
                                    )
                                    if image_part:
                                        image_bytes = image_part.blob
                                        # Determine image format
                                        content_type = image_part.content_type
                                        extension = self._get_image_extension(
                                            content_type
                                        )

                                        # Create image path
                                        image_path = f"{report_dir}/section_{section_num}_image_{image_index}.{extension}"
                                        images.append((image_path, image_bytes))
                                        image_index += 1
        except Exception as e:
            logger.warning(f"Warning: Error extracting images from paragraph: {e}")

        return images

    def _get_image_extension(self, content_type: str) -> str:
        """
        Get file extension from content type.

        Args:
            content_type: MIME type of the image

        Returns:
            File extension (without dot)
        """
        type_map = {
            "image/png": "png",
            "image/jpeg": "jpg",
            "image/jpg": "jpg",
            "image/gif": "gif",
            "image/bmp": "bmp",
            "image/tiff": "tiff",
            "image/svg+xml": "svg",
        }
        return type_map.get(content_type.lower(), "png")

    def _extract_table_data(self, table: Table) -> Dict[str, any]:
        """
        Extract structured data from a Word table.

        Args:
            table: Word table object

        Returns:
            Dictionary containing table data and metadata
        """
        try:
            rows_data = []

            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                rows_data.append(row_data)

            # Determine if first row is header
            has_header = False
            if len(rows_data) > 1:
                # Simple heuristic: check if first row cells are different from second row
                has_header = True  # Assume first row is header by default

            return {
                "rows": rows_data,
                "num_rows": len(rows_data),
                "num_cols": len(rows_data[0]) if rows_data else 0,
                "has_header": has_header,
                "header": rows_data[0] if has_header and rows_data else None,
                "data": (
                    rows_data[1:] if has_header and len(rows_data) > 1 else rows_data
                ),
            }

        except Exception as e:
            logger.warning(f"Warning: Error extracting table data: {e}")
            return {}

    def extract_text_by_paragraphs(self, docx_bytes: bytes) -> List[str]:
        """
        Extract text from Word document as a list of paragraphs.

        Args:
            docx_bytes: Word document content as bytes

        Returns:
            List of paragraph texts
        """
        try:
            docx_file = io.BytesIO(docx_bytes)
            document = Document(docx_file)

            paragraphs = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            return paragraphs

        except Exception as e:
            raise ValueError(f"Error extracting paragraphs: {e}")

    def extract_all_tables(self, docx_bytes: bytes) -> List[Dict[str, any]]:
        """
        Extract all tables from Word document.

        Args:
            docx_bytes: Word document content as bytes

        Returns:
            List of table data dictionaries
        """
        try:
            docx_file = io.BytesIO(docx_bytes)
            document = Document(docx_file)

            tables = []
            for table in document.tables:
                table_data = self._extract_table_data(table)
                if table_data:
                    tables.append(table_data)

            return tables

        except Exception as e:
            raise ValueError(f"Error extracting tables: {e}")

    def table_to_markdown(self, table_data: Dict[str, any]) -> str:
        """
        Convert table data to Markdown format.

        Args:
            table_data: Table data dictionary from _extract_table_data

        Returns:
            Markdown formatted table string
        """
        if not table_data or not table_data.get("rows"):
            return ""

        rows = table_data["rows"]
        if not rows:
            return ""

        # Build markdown table
        md_lines = []

        # Add header row
        if table_data.get("has_header") and table_data.get("header"):
            header = table_data["header"]
            md_lines.append("| " + " | ".join(header) + " |")
            # Add separator
            md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
            # Add data rows
            for row in table_data.get("data", []):
                md_lines.append("| " + " | ".join(row) + " |")
        else:
            # No header, treat all as data
            for row in rows:
                md_lines.append("| " + " | ".join(row) + " |")

        return "\n".join(md_lines)

    def table_to_csv(self, table_data: Dict[str, any], delimiter: str = ",") -> str:
        """
        Convert table data to CSV format.

        Args:
            table_data: Table data dictionary from _extract_table_data
            delimiter: CSV delimiter (default: comma)

        Returns:
            CSV formatted table string
        """
        if not table_data or not table_data.get("rows"):
            return ""

        rows = table_data["rows"]
        csv_lines = []

        for row in rows:
            # Escape cells that contain delimiter or quotes
            escaped_row = []
            for cell in row:
                if delimiter in cell or '"' in cell or "\n" in cell:
                    escaped_cell = '"' + cell.replace('"', '""') + '"'
                    escaped_row.append(escaped_cell)
                else:
                    escaped_row.append(cell)
            csv_lines.append(delimiter.join(escaped_row))

        return "\n".join(csv_lines)

    def extract_document_metadata(self, docx_bytes: bytes) -> Dict[str, any]:
        """
        Extract metadata from Word document.

        Args:
            docx_bytes: Word document content as bytes

        Returns:
            Dictionary containing document metadata
        """
        try:
            docx_file = io.BytesIO(docx_bytes)
            document = Document(docx_file)

            core_props = document.core_properties

            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": (
                    core_props.created.isoformat() if core_props.created else None
                ),
                "modified": (
                    core_props.modified.isoformat() if core_props.modified else None
                ),
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision or 0,
                "num_paragraphs": len(document.paragraphs),
                "num_tables": len(document.tables),
                "num_sections": len(document.sections),
            }

            return metadata

        except Exception as e:
            logger.warning(f"Warning: Error extracting metadata: {e}")
            return {}

    def chunk_text(self, text: str, words_per_chunk: int = 250) -> List[str]:
        """
        Split text into chunks of approximately the specified number of words.

        Args:
            text: Input text to chunk
            words_per_chunk: Target number of words per chunk

        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []

        words = text.split()
        chunks = []

        for i in range(0, len(words), words_per_chunk):
            chunk = " ".join(words[i : i + words_per_chunk])
            chunks.append(chunk)

        return chunks

    def clean_text(self, text: str) -> str:
        """
        Clean and normalize text content.

        Args:
            text: Input text to clean

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Basic text cleaning
        text = text.strip()
        # Remove excessive whitespace
        text = " ".join(text.split())

        return text

    def extract_chapters_from_markdown(self, docx_file_bytes):
        """Extract chapters and sections from DOCX by converting to markdown first"""
        try:
            # Step 1: Convert DOCX to markdown using MarkItDown
            md = MarkItDown()
            # Save DOCX bytes to temporary file
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
                temp_docx.write(docx_file_bytes)
                temp_docx_path = temp_docx.name

            try:
                # Convert to markdown
                result = md.convert(temp_docx_path)
                markdown_content = result.text_content

                # Step 2: Remove table of contents from markdown
                lines = markdown_content.split("\n")
                content_start = None
                for i, line in enumerate(lines):
                    line_stripped = line.strip()
                    # Look for the first actual chapter heading (not a TOC link)
                    if (
                        content_start is None
                        and re.match(r"^\*\*Chapter \d+", line_stripped)
                        and not re.search(r"\]\(#", line_stripped)
                    ):
                        content_start = i
                        break

                footnotes_start = None
                for i, line in enumerate(lines):
                    line_stripped = line.strip()

                    # Look for footnotes section start assuming the pattern is md standard
                    if "footnote-ref-1" in line_stripped and footnotes_start is None:
                        footnotes_start = i
                        logger.info(f"Found footnotes section starting at line {i}")
                        break
                # Separate content and footnotes
                if footnotes_start is not None:
                    footnote_lines = lines[footnotes_start:]
                    if content_start is not None:
                        clean_lines = lines[content_start:footnotes_start]
                    else:
                        clean_lines = lines[:footnotes_start]
                    logger.info(
                        f"Separated {len(footnote_lines)} footnote lines from content"
                    )

                else:
                    footnote_lines = []
                    if content_start is not None:
                        clean_lines = lines[content_start:]
                    else:
                        clean_lines = lines

                if footnote_lines:
                    footnote_lines = "\n".join(footnote_lines)
                    # Split footnotes by the pattern (#footnote-ref-N)
                    footnote_lines = re.split(r"\(#footnote-ref-\d+\)", footnote_lines)
                    # Filter out empty strings and strip whitespace
                    footnote_lines = [fn.strip() for fn in footnote_lines if fn.strip()]
                if content_start is not None:
                    logger.info(
                        f"Removed {content_start} lines of TOC, processing {len(clean_lines)} lines of content"
                    )
                else:
                    logger.info("No TOC detected, processing all lines")

                # Step 3: Extract chapters and sections from cleaned markdown
                chapters = {}
                current_chapter = None
                current_section = None
                current_content = []

                for i, line in enumerate(clean_lines):
                    line_stripped = line.strip()

                    # Skip empty lines but preserve them in content
                    if not line_stripped:
                        current_content.append(line)
                        continue

                    # Preserve bullet points and list items
                    if line_stripped.startswith(("- ", "* ", "• ", "+ ")):
                        current_content.append(line)
                        continue

                    # Preserve numbered lists (but not section numbers)
                    if re.match(
                        r"^\d+\.\s+[A-Z][a-z]", line_stripped
                    ):  # "1. Something" but not "1.1"
                        current_content.append(line)
                        continue

                    # Check for chapter headings (handle both regular dash and em dash)
                    chapter_match = re.match(
                        r"^\*\*Chapter (\d+)\s*[–—-]\s*(.+?)\*\*$", line_stripped
                    )
                    if chapter_match:
                        # Save previous section content if exists
                        if current_section and current_content:
                            if current_chapter not in chapters:
                                chapters[current_chapter] = {}
                            if footnote_lines:
                                for fn in re.findall(
                                    r"footnote-\d+", "\n".join(current_content)
                                ):
                                    fn_idx = int(fn.split("-")[-1]) - 1
                                    if 0 <= fn_idx < len(footnote_lines):
                                        current_content.append(
                                            f"[Footnote {fn_idx + 1}] : {footnote_lines[fn_idx]}]"
                                        )
                            chapters[current_chapter][current_section] = "\n".join(
                                current_content
                            ).strip()
                            current_content = []

                        chapter_num = int(chapter_match.group(1))
                        chapter_title = chapter_match.group(2).strip()
                        current_chapter = f"Chapter {chapter_num}"
                        current_section = None

                        if current_chapter not in chapters:
                            chapters[current_chapter] = {}

                        logger.info(f"Found {current_chapter}: {chapter_title}")
                        continue

                    # Check for section headings (e.g., **1.1 - Purpose**) (handle both regular dash and em dash)
                    section_match = re.match(
                        r"^\*\*(\d+\.\d+(?:\.\d+)*)\s*[–—-]\s*(.+?)\*\*$", line_stripped
                    )
                    if section_match:
                        # Save previous section content if exists
                        if current_section and current_content:
                            if current_chapter not in chapters:
                                chapters[current_chapter] = {}
                            if footnote_lines:
                                for fn in re.findall(
                                    r"footnote-\d+", "\n".join(current_content)
                                ):
                                    fn_idx = int(fn.split("-")[-1]) - 1
                                    if 0 <= fn_idx < len(footnote_lines):
                                        current_content.append(
                                            f"[Footnote {fn_idx + 1}] : {footnote_lines[fn_idx]}]"
                                        )
                            chapters[current_chapter][current_section] = "\n".join(
                                current_content
                            ).strip()
                            current_content = []

                        section_num = section_match.group(1)
                        section_title = section_match.group(2).strip()
                        current_section = section_num

                        # If no current chapter, derive from section number
                        if not current_chapter:
                            chapter_from_section = (
                                f"Chapter {section_num.split('.')[0]}"
                            )
                            current_chapter = chapter_from_section
                            if current_chapter not in chapters:
                                chapters[current_chapter] = {}
                            logger.info(
                                f"Determined chapter from section {current_chapter}: {chapter_title}"
                            )

                        logger.info(f"Found Section {section_num}: {section_title}")
                        continue

                    # Add content to current section
                    current_content.append(line)

                # Save final section content
                if current_section and current_content:
                    if current_chapter not in chapters:
                        chapters[current_chapter] = {}
                    if footnote_lines:
                        for fn in re.findall(
                            r"footnote-\d+", "\n".join(current_content)
                        ):
                            fn_idx = int(fn.split("-")[-1]) - 1
                            if 0 <= fn_idx < len(footnote_lines):
                                current_content.append(
                                    f"[Footnote {fn_idx + 1}] : {footnote_lines[fn_idx]}]"
                                )
                    chapters[current_chapter][current_section] = "\n".join(
                        current_content
                    ).strip()

                logger.info(f"Extracted {len(chapters)} chapters from markdown")

                return chapters

            finally:
                # Clean up temporary file
                if os.path.exists(temp_docx_path):
                    os.unlink(temp_docx_path)

        except Exception as e:
            logger.error(f"Error extracting chapters from markdown: {str(e)}")
            raise

    def extract_chapters_from_docx(self, docx_file_bytes):
        """Extract chapters and sections from DOCX file with their content"""
        try:
            import re

            doc = Document(io.BytesIO(docx_file_bytes))
            chapters = {}
            current_chapter = None
            current_section = None
            chapter_content = []

            # Regex patterns for chapter numbering - more restrictive to avoid false matches
            chapter_pattern = re.compile(
                r"^Chapter\s+(\d+)\s*[-–]?\s*(.*)$", re.IGNORECASE
            )
            main_chapter_pattern = re.compile(r"^(\d+)\.\s+(.+)$")
            section_pattern = re.compile(r"^(\d+\.\d+)\s+(.+)$")
            sub_section_pattern = re.compile(r"^(\d+\.\d+\.\d+)\s+(.+)$")

            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()

                if not text:
                    continue

                # Skip table of contents entries
                # TOC entries are typically in first ~200 paragraphs and have 'toc 2' style
                # or end with page numbers (indicating TOC entries)
                style_name = paragraph.style.name if paragraph.style else "No Style"
                has_page_ref = bool(
                    re.search(r"\t\d+$", text)
                )  # Ends with tab and page number
                is_likely_toc = i < 200 and (style_name == "toc 2" or has_page_ref)

                if is_likely_toc:
                    logger.debug("found TOC")
                    continue

                # Check if this is a list item (bullet or numbered)
                is_list_item = style_name and (
                    "list" in style_name.lower() or "bullet" in style_name.lower()
                )
                if is_list_item:
                    # Preserve list formatting with bullet symbol
                    chapter_content.append(f"• {text}")
                    continue

                # Check for Chapter (e.g., "Chapter 1 - Introduction")
                chapter_match = chapter_pattern.match(text)
                if chapter_match:
                    # Save previous chapter content
                    if current_chapter and chapter_content:
                        if current_chapter not in chapters:
                            chapters[current_chapter] = {}
                        if current_section:
                            chapters[current_chapter][current_section] = "\n".join(
                                chapter_content
                            )
                        else:
                            chapters[current_chapter]["content"] = "\n".join(
                                chapter_content
                            )

                    # Start new chapter
                    current_chapter = chapter_match.group(1)
                    current_section = None
                    chapter_content = []
                    continue

                # Check for main chapter (e.g., "1 Introduction")
                main_match = main_chapter_pattern.match(text)
                if main_match:
                    # Save previous chapter content
                    if current_chapter and chapter_content:
                        if current_chapter not in chapters:
                            chapters[current_chapter] = {}
                        if current_section:
                            chapters[current_chapter][current_section] = "\n".join(
                                chapter_content
                            )
                        else:
                            chapters[current_chapter]["content"] = "\n".join(
                                chapter_content
                            )

                    # Start new chapter
                    current_chapter = main_match.group(1)
                    current_section = None
                    chapter_content = []
                    continue

                # Check for sub-section (e.g., "1.3.1 Sound Banking")
                sub_sub_match = sub_section_pattern.match(text)
                if sub_sub_match:
                    section_number = sub_sub_match.group(1)
                    chapter_from_section = section_number.split(".")[0]

                    # Save previous section content if we have one
                    if current_chapter and current_section and chapter_content:
                        if current_chapter not in chapters:
                            chapters[current_chapter] = {}
                        chapters[current_chapter][current_section] = "\n".join(
                            chapter_content
                        )

                    # Only switch to the chapter if it exists in our chapters dict
                    # Otherwise, keep the current chapter context
                    if chapter_from_section in chapters:
                        current_chapter = chapter_from_section
                    elif current_chapter is None:
                        # If no current chapter, create the chapter
                        current_chapter = chapter_from_section
                        chapters[current_chapter] = {}

                    current_section = section_number
                    chapter_content = []
                    continue

                # Check for section (e.g., "1.1 Overview")
                sub_match = section_pattern.match(text)
                if sub_match:
                    section_number = sub_match.group(1)
                    chapter_from_section = section_number.split(".")[0]

                    # Save previous section content if we have one
                    if current_chapter and current_section and chapter_content:
                        if current_chapter not in chapters:
                            chapters[current_chapter] = {}
                        chapters[current_chapter][current_section] = "\n".join(
                            chapter_content
                        )

                    # Only switch to the chapter if it exists in our chapters dict
                    # Otherwise, keep the current chapter context
                    if chapter_from_section in chapters:
                        current_chapter = chapter_from_section
                    elif current_chapter is None:
                        # If no current chapter, create the chapter
                        current_chapter = chapter_from_section
                        chapters[current_chapter] = {}

                    current_section = section_number
                    chapter_content = []
                    continue

                # Add content to current chapter/section
                if current_chapter:
                    chapter_content.append(text)

            # Save the last chapter content
            if current_chapter and chapter_content:
                if current_chapter not in chapters:
                    chapters[current_chapter] = {}
                if current_section:
                    chapters[current_chapter][current_section] = "\n".join(
                        chapter_content
                    )
                else:
                    chapters[current_chapter]["content"] = "\n".join(chapter_content)

            logger.info(f"Extracted {len(chapters)} main chapters from DOCX")

            return chapters

        except Exception as e:
            logger.error(f"Error extracting chapters from DOCX: {str(e)}")
            raise

    def _debug_chapter_counts(self, chapters):
        """Debug method to count and log chapter statistics"""
        try:
            total_chapters = len(chapters)
            total_sections = 0

            debug_info = []
            debug_info.append("=== CHAPTER EXTRACTION DEBUG ===")
            debug_info.append(f"Total Chapters Found: {total_chapters}")

            for chapter_num, chapter_data in chapters.items():
                section_count = len(
                    [key for key in chapter_data.keys() if key != "content"]
                )
                total_sections += section_count
                debug_info.append(f"Chapter {chapter_num}: {section_count} sections")

            debug_info.append(f"Total Sections: {total_sections}")

            # Write debug info to file
            debug_path = os.path.join(os.getcwd(), "chapter_debug.txt")
            with open(debug_path, "w", encoding="utf-8") as f:
                f.write("\n".join(debug_info))

            logger.debug("\n".join(debug_info))

        except Exception as e:
            logger.debug(f"Debug error: {str(e)}")

    def _log_chapters_to_file(self, chapters):
        """Log extracted chapters to document_structure.md file for debugging"""
        try:
            from datetime import datetime

            # Create the markdown content
            md_content = ["# Document Structure - Extracted Chapters\n"]
            md_content.append(
                f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            )

            for main_chapter, content in chapters.items():
                md_content.append(f"## Chapter {main_chapter}\n")

                if isinstance(content, dict):
                    # Process sections
                    for section, section_content in content.items():
                        if section == "content":
                            # Main chapter content
                            if section_content:
                                md_content.append("### Main Content\n")
                                md_content.append(f"{section_content}\n\n")
                        else:
                            # Section content
                            if section_content:
                                md_content.append(f"### Section {section}\n")
                                md_content.append(f"{section_content}\n\n")
                else:
                    # Direct content for main chapter
                    if content:
                        md_content.append(f"{content}\n\n")

                md_content.append("---\n\n")

            # Write to file in the project root
            file_path = os.path.join(os.getcwd(), "document_structure.md")
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("".join(md_content))

        except Exception as e:
            logger.warning(f"Warning: Could not log chapters to file: {str(e)}")
            # Don't raise - this is just for debugging, shouldn't break the main flow

    def chapter_based_chunking(self, chapters_dict):
        """Create chunks based on chapters and sections with token limit handling

        Expected structure:
        - Chapters are main organizational units (Chapter 1, Chapter 2, etc.)
        - Sections are content units for chunking (1.1, 1.2, 2.1, 2.2, etc.)
        - Sub-sections are included with their parent sections (1.1.1, 1.1.2, etc.)
        """
        try:
            chunks = []
            max_tokens = 6000  # Safe limit below 8192 token max

            def estimate_tokens(text):
                # Rough estimation: ~4 characters per token
                return len(text) // 4

            def split_large_content(content, section_number, chapter, chunk_type):
                """Split large content into smaller chunks if it exceeds token limit"""
                if estimate_tokens(content) <= max_tokens:
                    return [
                        {
                            "section_number": section_number,
                            "content": content.strip(),
                            "chapter": chapter,
                            "chunk_type": chunk_type,
                        }
                    ]

                # Use existing chunk_text method to split large content
                # Adjust words_per_chunk to stay under token limit
                words_per_chunk = max_tokens // 6  # Conservative estimate
                text_chunks = self.chunk_text(content, words_per_chunk=words_per_chunk)

                result_chunks = []
                for text_chunk in text_chunks:
                    result_chunks.append(
                        {
                            "section_number": f"{section_number}",
                            "content": text_chunk.strip(),
                            "chapter": chapter,
                            "chunk_type": f"{chunk_type}_split",
                        }
                    )

                return result_chunks

            # Process the chapters dictionary
            for chapter_key, chapter_content in chapters_dict.items():
                # Determine if this is a chapter or a section
                if chapter_key.startswith("Chapter"):
                    chapter_name = chapter_key

                    if isinstance(chapter_content, dict):
                        # Process all sections within this chapter
                        for section_key, section_content in chapter_content.items():
                            if section_key != "content" and section_content:
                                # This is a section (e.g., "1.1", "1.2") within the chapter
                                split_chunks = split_large_content(
                                    section_content,
                                    section_key,
                                    chapter_name,
                                    "section",
                                )
                                chunks.extend(split_chunks)

                        # Process chapter-level content if exists
                        if "content" in chapter_content and chapter_content["content"]:
                            split_chunks = split_large_content(
                                chapter_content["content"],
                                chapter_name,
                                chapter_name,
                                "chapter_intro",
                            )
                            chunks.extend(split_chunks)
                    else:
                        # Direct content for chapter
                        if chapter_content:
                            split_chunks = split_large_content(
                                chapter_content,
                                chapter_name,
                                chapter_name,
                                "chapter_content",
                            )
                            chunks.extend(split_chunks)
                else:
                    # This is a chapter number (like "1", "2") from main_chapter_pattern match
                    chapter_name = f"Chapter {chapter_key}"

                    if isinstance(chapter_content, dict):
                        # Process all sections within this chapter individually
                        for section_key, section_content in chapter_content.items():
                            if section_key != "content" and section_content:
                                # Each section gets its own chunk with proper section number
                                split_chunks = split_large_content(
                                    section_content,
                                    section_key,  # Use actual section number (2.1, 2.2.1, etc.)
                                    chapter_name,
                                    "section",
                                )
                                chunks.extend(split_chunks)

                        # Process chapter-level content if exists
                        if "content" in chapter_content and chapter_content["content"]:
                            split_chunks = split_large_content(
                                chapter_content["content"],
                                chapter_key,  # Use chapter number as section for intro content
                                chapter_name,
                                "chapter_intro",
                            )
                            chunks.extend(split_chunks)
                    else:
                        # Direct content for chapter
                        if chapter_content:
                            split_chunks = split_large_content(
                                chapter_content, chapter_key, chapter_name, "section"
                            )
                            chunks.extend(split_chunks)

            # Filter out empty chunks
            chunks = [chunk for chunk in chunks if chunk["content"]]

            return chunks

        except Exception as e:
            logger.error(f"Error creating section-based chunks: {str(e)}")
            raise
