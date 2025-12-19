from manager.rag_orchestration import RAGOrchestrator
from config import Config
from fastapi import (
    FastAPI,
    UploadFile,
    File,
    Body,
    Request,
    status,
    HTTPException,
    WebSocket,
    WebSocketDisconnect,
)
from fastapi.openapi.utils import get_openapi
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import Optional, List, Dict, Any
import uuid
from datetime import datetime, timezone
import logging
import os
from pathlib import Path
import json
import aiofiles
import csv
import io
import base64
import re
import pandas as pd
from io import BytesIO, StringIO
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Protection


# Import our new upload system components
from manager.task_based_processor import FileProcessor


# Import chat history and session management
from service.chat_history import ChatHistoryService
from service.session_share import SessionShareService

# Import Azure AI Search Service for file listing
from service.azure_ai_search import AzureAISearchService

# Import JWT authentication services
from service.auth import initialize_jwt_service
from service.middleware import JWTAuthMiddleware, get_current_user_from_request
from service.utils import sanitize_filename

# Import content endpoints router
from content_endpoints import router as content_router
from content_endpoints import set_dependencies

# Import all models from the centralized data_model package
from model import (
    ErrorResponse,
    UploadResponse,
    StatusResponse,
    HealthResponse,
    ChatHistoryQuery,
    FeedbackUpdateRequest,
    ChatHistoryApiResponse,
    SessionShareRequest,
    SessionShareResponse,
    SessionShareInfoResponse,
    ChatExportRequest,
)
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER
from azure.storage.blob import BlobServiceClient, ContentSettings
from azure.identity import DefaultAzureCredential

from dotenv import load_dotenv

# Load environment variables from .env file only locally
# load_dotenv()
TELEMETRY_CONFIGURED = os.getenv("TELEMETRY_CONFIGURED", "false").lower() == "true"
try:
    # Simple, test-validated configuration for Azure Monitor
    from azure.monitor.opentelemetry import configure_azure_monitor

    TELEMETRY_CONFIGURED = False
    if os.getenv("APPLICATIONINSIGHTS_CONNECTION_STRING"):
        configure_azure_monitor(
            connection_string=os.getenv("APPLICATIONINSIGHTS_CONNECTION_STRING")
        )
        TELEMETRY_CONFIGURED = True
except ImportError:
    TELEMETRY_CONFIGURED = False

# Configure base logging levels for backend (match frontend style)

logger = logging.getLogger("backend")
BACKEND_EXCEPTION_TAG = "BACKEND_EXCEPTION"
BACKEND_EXCEPTION_TAG = "BACKEND_EXCEPTION"
_log_level_name = os.getenv("LOG_LEVEL", "DEBUG").upper()
_log_level = getattr(logging, _log_level_name, logging.INFO)
logger.setLevel(_log_level)

if not logger.handlers:
    handler = logging.StreamHandler()
    handler.setLevel(_log_level)
    handler.setFormatter(
        logging.Formatter("%(asctime)s %(levelname)s [%(name)s] %(message)s")
    )
    logger.addHandler(handler)

_backend_log_level = _log_level

# Align other backend-related loggers to the configured level
logging.getLogger("main").setLevel(_backend_log_level)
logging.getLogger("server.task_based_processor").setLevel(_backend_log_level)
logging.getLogger("uvicorn").setLevel(_backend_log_level)
logging.getLogger("fastapi").setLevel(_backend_log_level)

# Reduce noise from verbose dependencies
logging.getLogger("opentelemetry").setLevel(logging.WARNING)
logging.getLogger("azure").setLevel(logging.WARNING)
logging.getLogger("azure.core").setLevel(logging.WARNING)
logging.getLogger("azure.identity").setLevel(logging.WARNING)

logger.info("Backend telemetry configured: %s", TELEMETRY_CONFIGURED)

# Constants
BAD_REQUEST = "Bad Request"
INTERNAL_SERVER_ERROR = "Internal Server Error"
ACCESS_DENIED = "Access denied"

# Initialize configuration
config = Config()

config.validate_config()

# Initialize services with configuration
chat_history_service = (
    ChatHistoryService(BASE_URL=config.chat_history_api_url)
    if config.chat_history_enabled
    else None
)

# Initialize session share service with CosmosDB support
session_share_service = SessionShareService(
    chat_history_service=chat_history_service,
    bot_id=config.bot_id,
)


# Simple authentication dependency
def get_current_user(request: Request):
    """Get current user from JWT token via middleware"""
    try:
        current_user = get_current_user_from_request(request)
        user_id = (
            current_user.get("user_id")
            or current_user.get("sub")
            or current_user.get("oid")
        )

        if user_id:
            return {"UserID": user_id, "bot_id": config.bot_id}
        return None
    except Exception as e:
        logger.warning(f"Failed to get current user: {e}")
        return None


# WebSocket Connection Manager for real-time status updates
class ConnectionManager:
    """Manages WebSocket connections for real-time status updates"""

    def __init__(self):
        self.active_connections: List[WebSocket] = []
        self.work_id_subscriptions: Dict[str, List[WebSocket]] = {}

    async def connect(self, websocket: WebSocket):
        """Accept a new WebSocket connection"""
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        """Remove a WebSocket connection"""
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)

        # Remove from work_id subscriptions
        for work_id in self.work_id_subscriptions.keys():
            if websocket in self.work_id_subscriptions[work_id]:
                self.work_id_subscriptions[work_id].remove(websocket)
                if not self.work_id_subscriptions[work_id]:
                    del self.work_id_subscriptions[work_id]

    def subscribe_to_work_id(self, websocket: WebSocket, work_id: str):
        """Subscribe a WebSocket to updates for a specific work_id"""
        if work_id not in self.work_id_subscriptions:
            self.work_id_subscriptions[work_id] = []

        if websocket not in self.work_id_subscriptions[work_id]:
            self.work_id_subscriptions[work_id].append(websocket)

        logger.debug(f"WebSocket subscribed to work_id: {work_id}")

    async def send_personal_message(self, message: dict, websocket: WebSocket):
        """Send a message to a specific WebSocket connection"""
        try:
            await websocket.send_text(json.dumps(message))
        except Exception as e:
            logger.error(f"Error sending message to WebSocket: {e}")
            self.disconnect(websocket)

    async def broadcast_to_work_id(self, work_id: str, message: dict):
        """Broadcast a message to all WebSockets subscribed to a work_id"""
        if work_id in self.work_id_subscriptions:
            disconnected_connections = []

            for websocket in self.work_id_subscriptions[work_id]:
                try:
                    await websocket.send_text(json.dumps(message))
                except Exception as e:
                    logger.error(
                        f"Error broadcasting to WebSocket for work_id {work_id}: {e}"
                    )
                    disconnected_connections.append(websocket)

            # Clean up disconnected connections
            for websocket in disconnected_connections:
                self.disconnect(websocket)

    async def broadcast_to_all(self, message: dict):
        """Broadcast a message to all active WebSocket connections"""
        disconnected_connections = []

        for websocket in self.active_connections:
            try:
                await websocket.send_text(json.dumps(message))
            except Exception as e:
                logger.error(f"Error broadcasting to all WebSockets: {e}")
                disconnected_connections.append(websocket)

        # Clean up disconnected connections
        for websocket in disconnected_connections:
            self.disconnect(websocket)


# Initialize WebSocket manager
connection_manager = ConnectionManager()


# Custom OpenAPI function to enforce version 3.0.2 for APIM compatibility
def custom_openapi():
    if app.openapi_schema:
        return app.openapi_schema
    openapi_schema = get_openapi(
        title=app.title,
        version=app.version,
        description=app.description,
        routes=app.routes,
    )
    # Force OpenAPI 3.0.2 instead of 3.1.0 for APIM compatibility
    openapi_schema["openapi"] = "3.0.2"

    app.openapi_schema = openapi_schema
    return app.openapi_schema


# Create FastAPI app
app = FastAPI(
    title="Bot in a Box",
    version="1.0.0",
    description="Endpoints for Bot in a Box pattern.",
)

# Override the default OpenAPI function
app.openapi = custom_openapi

# Add CORS middleware to allow all hosts
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "*"
    ],  # Allows all origins TODO: change later to allow APIM only in prod
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

# Initialize services
try:
    orchestrator = RAGOrchestrator(chat_history_service=chat_history_service)
except Exception as e:
    logger.error(f"Failed to initialize RAGOrchestrator: {e}")
    orchestrator = None

# Initialize JWT authentication service
try:
    jwt_auth_service = initialize_jwt_service(config)
except Exception as e:
    logger.error(f"Failed to initialize JWT authentication service: {e}")
    jwt_auth_service = None

# Add JWT authentication middleware
try:
    app.add_middleware(JWTAuthMiddleware, jwt_service=jwt_auth_service)
except Exception as e:
    logger.error(f"Failed to add JWT authentication middleware: {e}")

try:
    file_processor = FileProcessor(max_workers=1, connection_manager=connection_manager)
    file_processor.start()
except Exception as e:
    logger.error(f"Failed to initialize upload system: {e}")
    file_processor = None

# Set dependencies for content_endpoints to avoid circular import
try:
    set_dependencies(
        orchestrator,
        config,
        chat_history_service,
        file_processor,
        session_share_service,
    )
except Exception as e:
    logger.error(f"Failed to set content endpoint dependencies: {e}")

# Include content endpoints router (query, PDF, image)
app.include_router(content_router)


# --- Helper functions ---


def get_azure_search_service():
    """Create and return an AzureAISearchService instance using global config"""
    try:
        # AzureAISearchService reads endpoint and index_name from global config
        return AzureAISearchService()
    except Exception:
        logger.exception("Failed to create AzureAISearchService")
        return None


# --- Non-streaming transformation function ---


@app.get(
    "/v1/sessions/titles",
    responses={
        200: {"description": "Session titles retrieved successfully"},
        400: {"model": ErrorResponse, "description": "Bad Request"},
        401: {"model": ErrorResponse, "description": "Unauthorized"},
        403: {
            "model": ErrorResponse,
            "description": "Forbidden - Authentication required",
        },
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Get session titles for a user (Authentication required)",
    tags=["Chat History"],
)
async def get_user_session_titles(
    request: Request,
    after_timestamp: Optional[str] = None,
):
    """
    Get all session titles for a user.

    Args:
        bot_id: Bot identifier
        user_id: User identifier
        after_timestamp: Optional ISO timestamp to filter sessions created after this time

    Returns:
        JSON response with session titles map
    """
    # Get current user from middleware
    current_user = get_current_user_from_request(request)
    user_id = current_user.get("user_id") or current_user.get("userID")
    bot_id = config.bot_id

    location = str(request.url)
    timestamp = datetime.now(timezone.utc).isoformat()

    try:
        if not chat_history_service:
            logger.error("[ERROR] [GET SESSION TITLES] Chat history service not available")
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=location,
                    message="Chat history service not available",
                    timestamp=timestamp,
                ).model_dump(),
            )

        # Convert after_timestamp string to datetime if provided
        after_dt = None
        if after_timestamp:
            try:
                after_dt = datetime.fromisoformat(
                    after_timestamp.replace("Z", "+00:00")
                )
            except ValueError as e:
                logger.error(f"[ERROR] [GET SESSION TITLES] Invalid timestamp format: {e}")
                return JSONResponse(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    content=ErrorResponse(
                        status=400,
                        reason="Bad Request",
                        location=location,
                        message=f"Invalid timestamp format: {after_timestamp}",
                        timestamp=timestamp,
                    ).model_dump(),
                )

        # Call the chat history service
        result = chat_history_service.get_sessions_with_titles(
            bot_id=bot_id, user_id=user_id, after_timestamp=after_dt
        )

        if result["success"]:
            return JSONResponse(status_code=status.HTTP_200_OK, content=result["data"])
        else:
            logger.error(
                f"[ERROR] [GET SESSION TITLES] Failed to retrieve session titles: {result.get('error')}"
            )
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=location,
                    message=result.get("error", "Failed to retrieve session titles"),
                    timestamp=timestamp,
                ).model_dump(),
            )

    except Exception as ex:
        logger.error(f"[ERROR] [GET SESSION TITLES] Exception: {ex}")
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=location,
                message=str(ex),
                timestamp=timestamp,
            ).model_dump(),
        )


@app.post(
    "/v1/upload",
    responses={
        200: {"description": "File uploaded successfully"},
        400: {"model": ErrorResponse, "description": "Bad Request"},
        401: {"model": ErrorResponse, "description": "Unauthorized"},
        403: {
            "model": ErrorResponse,
            "description": "Forbidden - Authentication required",
        },
        413: {"model": ErrorResponse, "description": "File too large"},
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Upload a file for processing (Authentication required)",
    tags=["File Upload"],
)
async def upload_file(
    request: Request,
    file: UploadFile = File(
        ...,
        description="File to upload (PDF, DOCX, DOC for documents; CSV, XLSX for metadata)",
    ),
):
    # Get current user from middleware (admin access already validated by middleware)

    # TODO: Enable later for audit logs

    # current_user = get_current_user_from_request(request)

    # Extract bot_id from headers (fallback to config if not provided)
    bot_id = config.bot_id

    location = str(request.url)
    timestamp = datetime.now(timezone.utc).isoformat()

    try:
        # Validate file
        if not file.filename:
            logger.error("[ERROR] [UPLOAD ERROR] No filename provided")
            return JSONResponse(
                status_code=status.HTTP_400_BAD_REQUEST,
                content=ErrorResponse(
                    status=400,
                    reason="Bad Request",
                    location=location,
                    message="No filename provided",
                    timestamp=timestamp,
                ).model_dump(),
            )
        file.filename = sanitize_filename(file.filename)

        # Detect file type
        file_extension = Path(file.filename).suffix.lower()
        logger.debug(f"[DEBUG] [FILE TYPE] Detected file extension: {file_extension}")

        # Route based on file type
        is_metadata_file = file_extension in {".csv", ".xlsx", ".xls"}
        is_document_file = file_extension in {".pdf", ".docx", ".doc"}

        if not (is_metadata_file or is_document_file):
            error_message = f"Unsupported file type '{file_extension}'. Supported: .pdf, .docx, .doc (documents), .csv, .xlsx, .xls (metadata)"
            logger.error(f"[ERROR] [UPLOAD ERROR] {error_message}")
            return JSONResponse(
                status_code=status.HTTP_400_BAD_REQUEST,
                content=ErrorResponse(
                    status=400,
                    reason="Bad Request",
                    location=location,
                    message=error_message,
                    timestamp=timestamp,
                ).model_dump(),
            )

        # Route to appropriate handler
        if is_metadata_file:
            return await handle_metadata_upload(file, location, timestamp, bot_id)
        else:
            return await handle_document_upload(file, location, timestamp, bot_id)

    except Exception as ex:
        logger.error(f"[ERROR] [UPLOAD EXCEPTION] Exception in upload_file: {ex}")
        logger.error(f"[ERROR] [UPLOAD EXCEPTION] Exception type: {type(ex).__name__}")
        logger.error(
            f"[ERROR] [UPLOAD EXCEPTION] File: {file.filename if file and hasattr(file, 'filename') else 'unknown'}"
        )

        logger.exception("[ERROR] [STACK TRACE] Full traceback captured")

        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=location,
                message=str(ex),
                timestamp=timestamp,
            ).model_dump(),
        )


async def handle_metadata_upload(
    file: UploadFile, location: str, timestamp: str, bot_id: str
):
    """Handle metadata file upload (CSV/Excel)"""

    file_extension = Path(file.filename).suffix.lower()

    try:
        # Read file content
        logger.debug("[DEBUG] [FILE READ] Reading uploaded metadata file")
        file_content = await file.read()
        file_size = len(file_content)

        # Validate the metadata file structure

        cfg = config  # Use global config instance

        # Read the file based on type
        if file_extension in [".xlsx", ".xls"]:
            df = pd.read_excel(BytesIO(file_content), engine="openpyxl")
        else:
            # Try common encodings for CSV
            df = None
            for encoding in ["utf-8", "utf-8-sig", "latin1", "iso-8859-1", "cp1252"]:
                try:
                    df = pd.read_csv(StringIO(file_content.decode(encoding)))
                    break
                except UnicodeDecodeError:
                    continue

            if df is None:
                raise ValueError("Could not read CSV with common encodings")

        # Validate headers
        missing_headers = [h for h in cfg.required_headers if h not in df.columns]
        if missing_headers:
            error_msg = f"Metadata file is missing required headers: {', '.join(missing_headers)}. Available headers: {', '.join(df.columns)}"
            logger.error(f"[ERROR] [METADATA VALIDATION] {error_msg}")
            return JSONResponse(
                status_code=status.HTTP_400_BAD_REQUEST,
                content=ErrorResponse(
                    status=400,
                    reason="Bad Request",
                    location=location,
                    message=error_msg,
                    timestamp=timestamp,
                ).model_dump(),
            )

        # Sanitize file_name column to match sanitized document filenames
        if "file_name" in df.columns:
            df["file_name"] = df["file_name"].apply(sanitize_filename)

        # Store metadata in memory for file processing
        timestamp_str = datetime.now().strftime("%Y%m%d%H%M%S")
        if file_processor:
            file_processor.set_metadata(df, timestamp_str)
        else:
            logger.warning(
                "[WARNING] [METADATA WARNING] File processor not available, metadata not cached"
            )

        # Upload to blob storage for audit/debugging (optional)
        try:
            ext = ".xlsx" if file_extension in [".xlsx", ".xls"] else ".csv"
            target_filename = f"metadata_{timestamp_str}{ext}"

            if file_processor and file_processor.blob_service:
                file_processor.blob_service.upload_bytes(target_filename, file_content)
            else:
                logger.warning(
                    "[WARNING] [BLOB UPLOAD] Blob service not available, skipping audit upload"
                )

        except Exception as e:
            logger.warning(
                f"[WARNING] [BLOB UPLOAD WARNING] Failed to upload to blob storage (metadata still loaded): {e}"
            )

        # Return success response
        return JSONResponse(
            status_code=status.HTTP_200_OK,
            content={
                "success": True,
                "file_type": "metadata",
                "message": "Metadata file validated and loaded successfully",
                "filename": file.filename,
                "size": file_size,
                "rows": len(df),
                "columns": list(df.columns),
                "timestamp": timestamp_str,
                "loaded_in_memory": True,
            },
        )

    except Exception as e:
        logger.error(f"[ERROR] [METADATA ERROR] {str(e)}")
        return JSONResponse(
            status_code=status.HTTP_400_BAD_REQUEST,
            content=ErrorResponse(
                status=400,
                reason="Bad Request",
                location=location,
                message=f"Invalid metadata file: {str(e)}",
                timestamp=timestamp,
            ).model_dump(),
        )


async def handle_document_upload(
    file: UploadFile, location: str, timestamp: str, bot_id: str
):
    """Handle document file upload (PDF/Word)"""

    try:
        # Create temp directory if it doesn't exist
        temp_dir = "temp_uploads"
        logger.debug(f"[DEBUG] [TEMP DIR] Ensuring temp directory exists: {temp_dir}")
        os.makedirs(temp_dir, exist_ok=True)

        # Generate unique filename for temp storage
        work_id = str(uuid.uuid4())
        temp_filename = f"{work_id}_{file.filename}"
        temp_file_path = os.path.join(temp_dir, temp_filename)

        # Save file to temp directory
        logger.debug("[DEBUG] [FILE SAVE] Saving uploaded file to temporary location")
        try:
            async with aiofiles.open(temp_file_path, "wb") as buffer:
                content = await file.read()
                await buffer.write(content)

            file_size = os.path.getsize(temp_file_path)

        except Exception as e:
            logger.error(f"[ERROR] [FILE SAVE ERROR] Failed to save uploaded file: {e}")
            logger.error(f"[ERROR] [FILE SAVE ERROR] Exception type: {type(e).__name__}")
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=location,
                    message="Failed to save uploaded file",
                    timestamp=timestamp,
                ).model_dump(),
            )

        # Create upload record and queue for processing
        logger.debug(
            "[DEBUG] [TASK MANAGER] Creating upload record and queueing for processing"
        )
        try:
            actual_work_id = file_processor.create_upload_record(
                original_filename=file.filename,
                file_path=temp_file_path,
                file_size=file_size,
                bot_id=bot_id,
            )
        except Exception as e:
            logger.error(f"[ERROR] [TASK MANAGER ERROR] Failed to create upload record: {e}")
            logger.error(f"[ERROR] [TASK MANAGER ERROR] Exception type: {type(e).__name__}")
            # Clean up the temp file since we couldn't create the record
            try:
                os.remove(temp_file_path)
                logger.debug("[DEBUG] [CLEANUP] Removed temp file due to task manager error")
            except OSError as cleanup_err:
                logger.warning(
                    "%s upload.cleanup_temp_file_failed path=%s",
                    BACKEND_EXCEPTION_TAG,
                    temp_file_path,
                    exc_info=cleanup_err,
                )
            raise

        return UploadResponse(
            work_id=actual_work_id,
            message="File uploaded successfully and queued for processing",
            filename=file.filename,
            file_size=file_size,
        )

    except Exception as ex:
        logger.error(f"[ERROR] [UPLOAD EXCEPTION] Exception in upload_file: {ex}")
        logger.error(f"[ERROR] [UPLOAD EXCEPTION] Exception type: {type(ex).__name__}")
        logger.error(
            f"[ERROR] [UPLOAD EXCEPTION] File: {file.filename if file and hasattr(file, 'filename') else 'unknown'}"
        )

        logger.exception("[ERROR] [STACK TRACE] Full traceback captured")

        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=location,
                message=str(ex),
                timestamp=timestamp,
            ).model_dump(),
        )


@app.get(
    "/v1/status/{work_id}",
    response_model=StatusResponse,
    responses={
        200: {"model": StatusResponse, "description": "Status retrieved successfully"},
        404: {"model": ErrorResponse, "description": "Work ID not found"},
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Get processing status for uploaded file",
    tags=["File Status"],
)
async def get_upload_status(work_id: str, request: Request):
    # Get current user from middleware

    location = str(request.url)
    timestamp = datetime.now(timezone.utc).isoformat()

    logger.debug(f"[DEBUG] [STATUS REQUEST] Status check requested for work_id: {work_id}")

    try:
        # Get status from task manager
        logger.debug(f"[DEBUG] [TASK MANAGER] Querying task manager for work_id: {work_id}")
        upload_record = file_processor.get_upload_info(work_id)

        if not upload_record:
            logger.warning(
                f"[WARNING] [STATUS NOT FOUND] Work ID not found in database: {work_id}"
            )
            return JSONResponse(
                status_code=status.HTTP_404_NOT_FOUND,
                content=ErrorResponse(
                    status=404,
                    reason="Not Found",
                    location=location,
                    message=f"Work ID {work_id} not found",
                    timestamp=timestamp,
                ).model_dump(),
            )

        logger.debug(
            f"[DEBUG] [STATUS FOUND] Found record for work_id {work_id}: status={upload_record['status']}, progress={upload_record.get('progress_percentage', 0)}%"
        )

        return StatusResponse(
            work_id=upload_record["work_id"],
            status=upload_record["status"],
            progress_percentage=upload_record.get("progress_percentage", 0),
            original_filename=upload_record["original_filename"],
            created_at=upload_record["created_at"],
            updated_at=upload_record.get("completed_at")
            or upload_record[
                "created_at"
            ],  # Use completed_at if available, otherwise created_at
            error_message=upload_record.get("error_message"),
            completed_at=upload_record.get("completed_at"),
        )

    except Exception as ex:
        logger.error(
            f"[ERROR] [STATUS ERROR] Exception in get_upload_status for work_id {work_id}: {ex}"
        )
        logger.error(f"[ERROR] [STATUS ERROR] Exception type: {type(ex).__name__}")
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=location,
                message=str(ex),
                timestamp=timestamp,
            ).model_dump(),
        )


@app.get(
    "/v1/botids/{bot_id}/listfiles",
    response_model=Dict[str, Any],
    responses={
        200: {"description": "Files retrieved successfully from blob storage"},
        404: {"model": ErrorResponse, "description": "File list not found"},
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Get file list from blob storage JSON file",
    tags=["File Management"],
)
async def list_files(
    request: Request,
    bot_id: str,
):
    """Get file list from a JSON file stored in blob storage"""
    # Get current user from middleware (admin access already validated by middleware)

    # TODO: Enable for audit logs
    # current_user = get_current_user_from_request(request)

    location = str(request.url)
    timestamp = datetime.now(timezone.utc).isoformat()

    try:
        # Check if file processor and blob service are available
        if not file_processor or not file_processor.blob_service:
            logger.error("[ERROR] [LIST FILES] Blob service not available")
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=location,
                    message="Blob storage service not available",
                    timestamp=timestamp,
                ).model_dump(),
            )

        # Construct the blob file name
        file_list_name = f"{bot_id}-filelist.json"

        # Also check for files under the config bot_id (for migration purposes)
        config_file_list_name = f"{config.bot_id}-filelist.json"

        merged_files = []
        latest_metadata = {}

        try:
            # Try to download the file list from blob storage
            file_content = file_processor.blob_service.download_bytes(file_list_name)

            # Parse JSON content
            file_list_data = json.loads(file_content.decode("utf-8"))
            merged_files.extend(file_list_data.get("files", []))
            latest_metadata = file_list_data

        except Exception as download_ex:
            if "BlobNotFound" in str(download_ex) or "404" in str(download_ex):
                logger.info(f"[INFO] [LIST FILES] No file list found for {file_list_name}")
            else:
                logger.warning(
                    f"[WARNING] [LIST FILES] Error downloading {file_list_name}: {download_ex}"
                )

        # Also check config bot_id file (for files uploaded before the bot_id fix)
        if bot_id != config.bot_id:
            try:
                config_file_content = file_processor.blob_service.download_bytes(
                    config_file_list_name
                )
                config_file_list_data = json.loads(config_file_content.decode("utf-8"))
                config_files = config_file_list_data.get("files", [])

                # Add files from config that aren't already in the main list
                existing_filenames = {
                    f.get("name", f.get("file_name", "")) for f in merged_files
                }
                for config_file in config_files:
                    config_filename = config_file.get(
                        "name", config_file.get("file_name", "")
                    )
                    if config_filename and config_filename not in existing_filenames:
                        merged_files.append(config_file)

            except Exception as config_download_ex:
                if "BlobNotFound" in str(config_download_ex) or "404" in str(
                    config_download_ex
                ):
                    logger.info(
                        f"[INFO] [LIST FILES] No config file list found for {config_file_list_name}"
                    )
                else:
                    logger.warning(
                        f"[WARNING] [LIST FILES] Error downloading {config_file_list_name}: {config_download_ex}"
                    )

        if merged_files:
            # Create merged file list data
            merged_data = {
                "bot_id": bot_id,
                "updated_at": latest_metadata.get("updated_at", timestamp),
                "updated_by": latest_metadata.get("updated_by", "system"),
                "total_files": len(merged_files),
                "files": merged_files,
                "merged_from_config": bot_id != config.bot_id
                and len(merged_files) > len(latest_metadata.get("files", [])),
            }

            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "success": True,
                    "bot_id": bot_id,
                    "file_list": merged_data,
                    "timestamp": timestamp,
                },
            )
        else:
            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "success": True,
                    "bot_id": bot_id,
                    "file_list": {
                        "bot_id": bot_id,
                        "updated_at": timestamp,
                        "updated_by": "system",
                        "total_files": 0,
                        "files": [],
                    },
                    "timestamp": timestamp,
                    "message": "No file list found, returning empty list",
                },
            )

    except Exception as ex:
        logger.error(f"[ERROR] [LIST FILES] Exception in list_files: {ex}")
        logger.error(f"[ERROR] [LIST FILES] Exception type: {type(ex).__name__}")
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=location,
                message=f"Failed to retrieve file list: {str(ex)}",
                timestamp=timestamp,
            ).model_dump(),
        )


@app.patch(
    "/v1/botids/{bot_id}/updatefilelist",
    response_model=Dict[str, Any],
    responses={
        200: {"description": "File list updated successfully"},
        400: {"model": ErrorResponse, "description": "Invalid JSON data"},
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Update or create file list in blob storage",
    tags=["File Management"],
)
async def update_file_list(
    request: Request,
    bot_id: str,
    file_list_data: Dict[str, Any],
):
    """Update or create a file list JSON file in blob storage"""
    # Get current user from middleware (admin access already validated by middleware)
    current_user = get_current_user_from_request(request)

    location = str(request.url)
    timestamp = datetime.now(timezone.utc).isoformat()

    try:
        # Check if file processor and blob service are available
        if not file_processor or not file_processor.blob_service:
            logger.error("[ERROR] [UPDATE FILE LIST] Blob service not available")
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=location,
                    message="Blob storage service not available",
                    timestamp=timestamp,
                ).model_dump(),
            )

        # Construct the blob file name
        file_list_name = f"{bot_id}-filelist.json"

        # Add metadata to the file list data
        updated_data = {
            "bot_id": bot_id,
            "updated_at": timestamp,
            "updated_by": (
                current_user.get("sub", "unknown") if current_user else "unknown"
            ),
            "files": file_list_data.get("files", []),
            **{
                k: v
                for k, v in file_list_data.items()
                if k not in ["bot_id", "updated_at", "updated_by"]
            },
        }

        # Convert to JSON bytes
        json_content = json.dumps(updated_data, indent=2).encode("utf-8")

        # Upload to blob storage
        file_processor.blob_service.upload_bytes(file_list_name, json_content)

        return JSONResponse(
            status_code=status.HTTP_200_OK,
            content={
                "success": True,
                "bot_id": bot_id,
                "message": f"File list updated successfully for bot {bot_id}",
                "file_name": file_list_name,
                "updated_at": timestamp,
                "file_count": len(updated_data.get("files", [])),
                "timestamp": timestamp,
            },
        )

    except json.JSONDecodeError as json_ex:
        logger.error(f"[ERROR] [UPDATE FILE LIST] Invalid JSON data: {json_ex}")
        return JSONResponse(
            status_code=status.HTTP_400_BAD_REQUEST,
            content=ErrorResponse(
                status=400,
                reason="Bad Request",
                location=location,
                message=f"Invalid JSON data: {str(json_ex)}",
                timestamp=timestamp,
            ).model_dump(),
        )
    except Exception as ex:
        logger.error(f"[ERROR] [UPDATE FILE LIST] Exception in update_file_list: {ex}")
        logger.error(f"[ERROR] [UPDATE FILE LIST] Exception type: {type(ex).__name__}")
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=location,
                message=f"Failed to update file list: {str(ex)}",
                timestamp=timestamp,
            ).model_dump(),
        )


@app.delete(
    "/v1/files/{file_name}",
    responses={
        200: {"description": "File deleted successfully"},
        401: {"model": ErrorResponse, "description": "Unauthorized"},
        403: {
            "model": ErrorResponse,
            "description": "Forbidden - Admin access required",
        },
        404: {"model": ErrorResponse, "description": "File not found"},
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Delete a file from Azure AI Search and Blob Storage (Admin only)",
    tags=["File Management"],
)
async def delete_azure_file(file_name: str, request: Request):
    # Get current user from middleware (admin access already validated by middleware)
    # TODO: Enable later for audit logs
    # current_user = get_current_user_from_request(request)

    location = str(request.url)
    timestamp = datetime.now(timezone.utc).isoformat()

    try:
        # Get bot_id from config
        bot_id = config.bot_id

        # Track results from both operations
        search_result = None
        blob_result = None

        # 1. Delete from Azure AI Search
        search_service = get_azure_search_service()

        if search_service is None:
            logger.error(
                "[ERROR] [DELETE FILE] Failed to create Azure Search service instance"
            )
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=location,
                    message="Failed to initialize Azure Search service",
                    timestamp=timestamp,
                ).model_dump(),
            )

        search_result = search_service.delete_file_documents(file_name)

        # 2. Delete from Blob Storage and update file list
        if file_processor and file_processor.blob_service:
            blob_result = file_processor.blob_service.delete_file_and_update_list(
                file_name, bot_id
            )
        else:
            logger.error("[ERROR] [DELETE FILE] Blob service not available")
            blob_result = {
                "success": False,
                "message": "Blob service not available",
                "deleted_blob_count": 0,
            }

        # Determine overall success
        # Consider it successful if at least one operation succeeded
        overall_success = search_result["success"] or (
            blob_result and blob_result["success"]
        )

        if overall_success:
            # Build success message
            messages = []
            if search_result["success"]:
                messages.append(
                    f"Deleted {search_result['deleted_count']} document(s) from search index"
                )
            if blob_result and blob_result["success"]:
                messages.append(
                    f"Deleted {blob_result['deleted_blob_count']} blob(s) from storage"
                )

            success_message = "; ".join(messages)

            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "success": True,
                    "message": success_message,
                    "file_name": file_name,
                    "search_deleted_count": search_result.get("deleted_count", 0),
                    "blob_deleted_count": (
                        blob_result.get("deleted_blob_count", 0) if blob_result else 0
                    ),
                    "removed_from_list": (
                        blob_result.get("removed_from_list", False)
                        if blob_result
                        else False
                    ),
                    "timestamp": timestamp,
                },
            )
        else:
            # Both operations failed
            error_messages = []
            if not search_result["success"]:
                error_messages.append(f"Search: {search_result['message']}")
            if blob_result and not blob_result["success"]:
                error_messages.append(f"Storage: {blob_result['message']}")

            error_message = "; ".join(error_messages)

            logger.error(f"[ERROR] [DELETE FILE] Failed to delete file: {error_message}")
            return JSONResponse(
                status_code=status.HTTP_404_NOT_FOUND,
                content=ErrorResponse(
                    status=404,
                    reason="Not Found",
                    location=location,
                    message=f"Failed to delete file: {error_message}",
                    timestamp=timestamp,
                ).model_dump(),
            )

    except Exception as ex:
        logger.exception("[ERROR] [DELETE FILE] Exception in delete_azure_file")
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=location,
                message=f"Failed to delete file: {str(ex)}",
                timestamp=timestamp,
            ).model_dump(),
        )


@app.get(
    "/v1/health",
    response_model=HealthResponse,
    responses={
        200: {"model": HealthResponse, "description": "System health check passed"},
        503: {
            "model": ErrorResponse,
            "description": "Service unavailable - some components unhealthy",
        },
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Health check endpoint for system status",
    tags=["Health"],
)
async def health_check(request: Request):
    # Get current user from middleware
    # current_user = get_current_user_from_request(request)

    location = str(request.url)
    timestamp = datetime.now(timezone.utc).isoformat()

    services = {}
    overall_healthy = True

    try:
        # Check file processor and task manager
        try:
            if file_processor is None:
                logger.error(
                    "[ERROR] File Processor: Not initialized (file_processor is None)"
                )
                services["file_processor"] = {
                    "status": "unhealthy",
                    "message": "File processor not initialized",
                }
                overall_healthy = False
            else:
                # Get health status from the processor
                health_info = file_processor.check_service_health()

                # Add file processor status
                services["file_processor"] = {
                    "status": health_info.get("status", "unknown"),
                    "message": "File processor and task manager operational",
                }

                # Flatten and add individual services from the processor
                processor_services = health_info.get("services", {})
                for service_name, service_info in processor_services.items():
                    services[f"file_processor_{service_name}"] = service_info

                if health_info.get("status") != "healthy":
                    overall_healthy = False
        except Exception as fp_ex:
            logger.error(f"[ERROR] File Processor: Error during check - {str(fp_ex)}")
            logger.error(
                f"File processor error details: {type(fp_ex).__name__}: {str(fp_ex)}"
            )
            services["file_processor"] = {
                "status": "unhealthy",
                "message": f"File processor error: {str(fp_ex)}",
            }
            overall_healthy = False

        # Check model orchestrator
        try:
            if orchestrator is None:
                logger.error(
                    "[ERROR] Model Orchestrator: Not initialized (orchestrator is None)"
                )
                services["model_orchestrator"] = {
                    "status": "unhealthy",
                    "message": "Model orchestrator is not available",
                }
                overall_healthy = False
            else:
                services["model_orchestrator"] = {
                    "status": "healthy",
                    "message": "Model orchestrator is available",
                }
        except Exception as mo_ex:
            logger.error(f"[ERROR] Model Orchestrator: Error during check - {str(mo_ex)}")
            logger.error(
                f"Model orchestrator error details: {type(mo_ex).__name__}: {str(mo_ex)}"
            )
            services["model_orchestrator"] = {
                "status": "unhealthy",
                "message": f"Model orchestrator error: {str(mo_ex)}",
            }
            overall_healthy = False

        # Check Azure services through file processor
        try:
            if file_processor is None:
                logger.warning(
                    "[WARNING] Azure Services: Cannot check - file_processor is None"
                )
                services["azure_services"] = {
                    "status": "unknown",
                    "message": "Cannot check Azure services - file processor not available",
                }
            elif hasattr(file_processor, "check_service_health"):
                azure_health = file_processor.check_service_health()

                # Extract Azure services from the nested structure
                azure_services = azure_health.get("services", {})
                all_azure_healthy = True
                azure_issues = []

                for service_name in [
                    "azure_search",
                    "blob_storage",
                    "llm_client",
                ]:
                    if service_name in azure_services:
                        service_info = azure_services[service_name]
                        service_status = service_info.get("status", "unknown")
                        service_message = service_info.get("message", "")

                        if service_status == "healthy":
                            logger.info(f"[INFO] Azure {service_name}: {service_message}")
                        elif service_status == "unhealthy":
                            all_azure_healthy = False
                            azure_issues.append(f"{service_name}: {service_message}")
                            logger.error(f"[ERROR] Azure {service_name}: {service_message}")
                        else:
                            azure_issues.append(f"{service_name}: {service_status}")
                            logger.warning(f"[WARNING] Azure {service_name}: {service_status}")
                    else:
                        all_azure_healthy = False
                        azure_issues.append(
                            f"{service_name}: not found in health response"
                        )
                        logger.warning(
                            f"[WARNING] Azure {service_name}: not found in health response"
                        )

                if all_azure_healthy and not azure_issues:
                    services["azure_services"] = {
                        "status": "healthy",
                        "message": "All Azure services are initialized and operational",
                    }
                elif azure_issues:
                    # Check overall health status from the file processor
                    processor_status = azure_health.get("status", "unknown")
                    if processor_status == "healthy":
                        services["azure_services"] = {
                            "status": "degraded",
                            "message": f"Some Azure services have issues: {'; '.join(azure_issues)}",
                        }
                        logger.warning(
                            f"[WARNING] Azure Services: Some issues detected - {'; '.join(azure_issues)}"
                        )
                    else:
                        services["azure_services"] = {
                            "status": "unhealthy",
                            "message": f"Azure services unhealthy: {'; '.join(azure_issues)}",
                        }
                        logger.error(
                            f"[ERROR] Azure Services: Unhealthy - {'; '.join(azure_issues)}"
                        )
                        overall_healthy = False
                else:
                    services["azure_services"] = {
                        "status": "healthy",
                        "message": "Azure services are operational",
                    }
            else:
                logger.warning(
                    "[WARNING] Azure Services: check_service_health method not available"
                )
                services["azure_services"] = {
                    "status": "unknown",
                    "message": "Azure services health check not available",
                }
        except Exception as az_ex:
            logger.error(f"[ERROR] Azure Services: Error during check - {str(az_ex)}")
            logger.error(
                f"Azure services error details: {type(az_ex).__name__}: {str(az_ex)}"
            )
            services["azure_services"] = {
                "status": "unhealthy",
                "message": f"Azure services error: {str(az_ex)}",
            }
            overall_healthy = False

        if overall_healthy:
            logger.info("[INFO] Health Check PASSED: All critical services are healthy")
            return HealthResponse(
                status="healthy", timestamp=timestamp, services=services
            )
        else:
            # Count unhealthy services for summary
            unhealthy_services = [
                name for name, info in services.items() if info["status"] == "unhealthy"
            ]
            warning_services = [
                name for name, info in services.items() if info["status"] == "warning"
            ]

            logger.error(
                f"[ERROR] Health Check FAILED: {len(unhealthy_services)} services unhealthy"
            )
            if unhealthy_services:
                logger.error(f"Unhealthy services: {', '.join(unhealthy_services)}")
            if warning_services:
                logger.warning(f"Services with warnings: {', '.join(warning_services)}")

            return JSONResponse(
                status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                content=ErrorResponse(
                    status=503,
                    reason="Service Unavailable",
                    location=location,
                    message="One or more services are unhealthy",
                    timestamp=timestamp,
                ).model_dump(),
            )

    except Exception as ex:
        logger.error(f"EXCEPTION in health_check: {ex}")
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason=INTERNAL_SERVER_ERROR,
                location=location,
                message=f"Health check failed: {str(ex)}",
                timestamp=timestamp,
            ).model_dump(),
        )


@app.get(
    "/v1/config",
    response_model=Dict[str, Any],
    responses={
        200: {"description": "Configuration retrieved successfully"},
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Get bot configuration",
    tags=["Configuration"],
)
async def get_config(request: Request):
    """Get bot configuration settings for frontend"""
    try:
        cfg = config  # Use global config instance
        try:
            if cfg.azure_storage_account_name and cfg.azure_storage_container_name:
                credential = DefaultAzureCredential()
                account_url = (
                    f"https://{cfg.azure_storage_account_name}.blob.core.windows.net"
                )
                blob_service_client = BlobServiceClient(
                    account_url=account_url, credential=credential
                )
                container_client = blob_service_client.get_container_client(
                    cfg.azure_storage_container_name
                )
                config_blob_name = "config.json"
                blob_client = container_client.get_blob_client(config_blob_name)

                if blob_client.exists():
                    blob_data = blob_client.download_blob().readall()
                    bot_config = json.loads(blob_data.decode("utf-8"))

                    # --- Completely flatten look_and_feel into root and remove the field ---
                    if "look_and_feel" in bot_config and isinstance(
                        bot_config["look_and_feel"], dict
                    ):
                        lf = bot_config["look_and_feel"]

                        # Move all look_and_feel fields to root (root always gets overridden)
                        for k, v in lf.items():
                            bot_config[k] = v

                        # Remove the look_and_feel field entirely
                        del bot_config["look_and_feel"]

                    images: Dict[str, Optional[str]] = {
                        "logo_base64": None,
                        "bot_icon_base64": None,
                        "user_icon_base64": None,
                    }
                    fname_map = {
                        "logo_base64": bot_config.get("logo_file_name"),
                        "bot_icon_base64": bot_config.get("bot_icon_file_name"),
                        "user_icon_base64": bot_config.get("user_icon_file_name"),
                    }

                    for resp_key, fname in fname_map.items():
                        if not fname:
                            images[resp_key] = None
                            continue
                        try:
                            blob_bytes = (
                                container_client.get_blob_client(fname)
                                .download_blob()
                                .readall()
                            )
                            images[resp_key] = base64.b64encode(blob_bytes).decode(
                                "utf-8"
                            )
                        except Exception as img_ex:
                            logger.warning(
                                f"[WARNING] [/v1/config] Failed to read '{fname}': {img_ex}"
                            )
                            images[resp_key] = None

                    # Compose consistent config body
                    has_filters = bot_config.get("has_filters", cfg.has_filters)
                    response_config = {
                        "has_filters": has_filters,
                        "filters": (
                            bot_config.get("filters", cfg.filters)
                            if has_filters
                            else {}
                        ),
                        "filter_mapping": (
                            bot_config.get("filter_mapping", cfg.filter_mapping)
                            if has_filters
                            else {}
                        ),
                        "bot_name": bot_config.get(
                            "bot_name", cfg.bot_config.get("bot_name")
                        ),
                        "version": bot_config.get(
                            "version", cfg.bot_config.get("version")
                        ),
                        "language": bot_config.get(
                            "language", cfg.bot_config.get("language")
                        ),
                        "about_text": bot_config.get(
                            "about_text", cfg.bot_config.get("about_text")
                        ),
                        "disclaimer_text": bot_config.get(
                            "disclaimer_text", cfg.bot_config.get("disclaimer_text")
                        ),
                        "primary_color": bot_config.get(
                            "primary_color", cfg.bot_config.get("primary_color")
                        ),
                        "secondary_background_color": bot_config.get(
                            "secondary_background_color",
                            cfg.bot_config.get("secondary_background_color"),
                        ),
                        "background_color": bot_config.get(
                            "background_color", cfg.bot_config.get("background_color")
                        ),
                        "text_color": bot_config.get(
                            "text_color", cfg.bot_config.get("text_color")
                        ),
                        "font_family": bot_config.get(
                            "font_family", cfg.bot_config.get("font_family")
                        ),
                        "font_size": bot_config.get(
                            "font_size", cfg.bot_config.get("font_size")
                        ),
                        "welcome_message": bot_config.get(
                            "welcome_message", cfg.bot_config.get("welcome_message")
                        ),
                        "default_response": bot_config.get(
                            "default_response", cfg.bot_config.get("default_response")
                        ),
                        "feedback_contact_name": bot_config.get(
                            "feedback_contact_name",
                            cfg.bot_config.get("feedback_contact_name"),
                        ),
                        "feedback_contact_email": bot_config.get(
                            "feedback_contact_email",
                            cfg.bot_config.get("feedback_contact_email"),
                        ),
                        "faq": bot_config.get("faq", cfg.bot_config.get("faq", [])),
                        "system_prompt": bot_config.get(
                            "system_prompt", cfg.bot_config.get("system_prompt", "")
                        ),
                        "external_links": bot_config.get(
                            "external_links", cfg.bot_config.get("external_links", [])
                        ),
                        "images": images,
                    }

                    # Envelope to match PUT shape
                    return {
                        "status": "success",
                        "message": "Configuration retrieved successfully",
                        "config": response_config,
                        "updated_by": bot_config.get(
                            "updated_by"
                        ),  # may be None if not tracked
                        "timestamp": datetime.now(timezone.utc).isoformat(),
                    }

                else:
                    logger.warning(
                        f"Bot-specific config {config_blob_name} not found, using default config"
                    )

        except Exception as e:
            logger.error(f"Error loading bot-specific config: {e}")
            # Fall back to default config if bot-specific config fails

        # Build defaults path (no blob)
        # --- Completely flatten look_and_feel into root and remove the field ---
        if "look_and_feel" in cfg and isinstance(cfg["look_and_feel"], dict):
            lf = cfg["look_and_feel"]

            # Move all look_and_feel fields to root (root always gets overridden)
            for k, v in lf.items():
                cfg[k] = v

            # Remove the look_and_feel field entirely
            del cfg["look_and_feel"]
        images_default = {
            "logo_base64": None,
            "bot_icon_base64": None,
            "user_icon_base64": None,
        }
        has_filters = cfg.has_filters
        response_config = {
            "has_filters": has_filters,
            "filters": cfg.filters if has_filters else {},
            "filter_mapping": cfg.filter_mapping if has_filters else {},
            "bot_name": cfg.bot_config.get("bot_name"),
            "version": cfg.bot_config.get("version"),
            "language": cfg.bot_config.get("language"),
            "about_text": cfg.bot_config.get("about_text"),
            "disclaimer_text": cfg.bot_config.get("disclaimer_text"),
            "primary_color": cfg.bot_config.get("primary_color"),
            "secondary_background_color": cfg.bot_config.get(
                "secondary_background_color"
            ),
            "background_color": cfg.bot_config.get("background_color"),
            "text_color": cfg.bot_config.get("text_color"),
            "font_family": cfg.bot_config.get("font_family"),
            "font_size": cfg.bot_config.get("font_size"),
            "welcome_message": cfg.bot_config.get("welcome_message"),
            "default_response": cfg.bot_config.get("default_response"),
            "feedback_contact_name": cfg.bot_config.get("feedback_contact_name"),
            "feedback_contact_email": cfg.bot_config.get("feedback_contact_email"),
            "faq": cfg.bot_config.get("faq", []),
            "system_prompt": cfg.bot_config.get("system_prompt"),
            "external_links": cfg.bot_config.get("external_links", []),
            "images": images_default,
        }
        return {
            "status": "success",
            "message": "Configuration retrieved successfully",
            "config": response_config,
            "updated_by": None,
            "timestamp": datetime.now(timezone.utc).isoformat(),
        }

    except Exception as e:
        logger.error(f"Error loading config: {e}")
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=str(request.url),
                message=f"Failed to load configuration: {str(e)}",
                timestamp=datetime.now(timezone.utc).isoformat(),
            ).model_dump(),
        )


@app.delete(
    "/v1/reset-factory-new",
    responses={
        200: {"description": "Factory reset completed successfully"},
        401: {"model": ErrorResponse, "description": "Unauthorized"},
        403: {
            "model": ErrorResponse,
            "description": "Forbidden - Factory reset disabled",
        },
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Factory reset - Delete all files, images, search index, and reset config (SuperAdmin only)",
    tags=["Administration"],
)
async def factory_reset(request: Request):
    """
    Factory reset endpoint - DESTRUCTIVE OPERATION

    This will:
    1. Delete all uploaded files from blob storage
    2. Delete all images from blob storage
    3. Delete all documents from Azure AI Search index
    4. Reset config.json to default values

    Only available if FACTORY_RESET_BOT environment variable is set to "true"
    Requires SuperAdmin access
    """
    # Get current user from middleware (superadmin access already validated by middleware)
    current_user = get_current_user_from_request(request)

    location = str(request.url)
    timestamp = datetime.now(timezone.utc).isoformat()

    logger.warning("[WARNING] [FACTORY RESET] Factory reset requested!")
    logger.warning(
        f"[WARNING] [FACTORY RESET] Requested by user: {current_user.get('sub', 'unknown')}"
    )

    try:
        # Check if factory reset is enabled via environment variable
        factory_reset_enabled = (
            os.getenv("FACTORY_RESET_BOT", "false").lower() == "true"
        )

        if factory_reset_enabled:
            logger.warning(
                "[WARNING] [FACTORY RESET] Factory reset is ENABLED - proceeding with reset!"
            )

            reset_results = {
                "blob_storage": {"success": False, "message": "", "deleted_count": 0},
                "azure_search": {"success": False, "message": "", "deleted_count": 0},
                "config_reset": {"success": False, "message": ""},
            }

            # 1. Delete all files from blob storage
            logger.warning(
                "[WARNING] [FACTORY RESET] Step 1: Deleting all files from blob storage..."
            )
            try:
                if file_processor and file_processor.blob_service:
                    # List all blobs
                    deleted_count = file_processor.blob_service.delete_all_blobs()
                    reset_results["blob_storage"]["success"] = True
                    reset_results["blob_storage"]["deleted_count"] = deleted_count
                    reset_results["blob_storage"][
                        "message"
                    ] = f"Deleted {deleted_count} files from blob storage"
                    logger.warning(
                        f"[WARNING] [FACTORY RESET] Deleted {deleted_count} files from blob storage"
                    )
                else:
                    reset_results["blob_storage"][
                        "message"
                    ] = "Blob service not available"
                    logger.error("[ERROR] [FACTORY RESET] Blob service not available")
            except Exception as blob_ex:
                reset_results["blob_storage"]["message"] = f"Error: {str(blob_ex)}"
                logger.error(f"[ERROR] [FACTORY RESET] Blob deletion error: {blob_ex}")

            # 2. Delete all documents from Azure AI Search
            logger.warning(
                "[WARNING] [FACTORY RESET] Step 2: Deleting all documents from Azure AI Search..."
            )
            try:
                search_service = get_azure_search_service()
                if search_service:
                    # Delete all documents by searching for all and deleting
                    delete_result = search_service.delete_all_documents()

                    reset_results["azure_search"]["success"] = delete_result.get(
                        "success", False
                    )
                    reset_results["azure_search"]["deleted_count"] = delete_result.get(
                        "deleted_count", 0
                    )
                    reset_results["azure_search"]["message"] = delete_result.get(
                        "message", ""
                    )
                    logger.warning(
                        f"[WARNING] [FACTORY RESET] Azure Search: {delete_result.get('message')}"
                    )
                else:
                    reset_results["azure_search"][
                        "message"
                    ] = "Azure Search service not available"
                    logger.error(
                        "[ERROR] [FACTORY RESET] Azure Search service not available"
                    )
            except Exception as search_ex:
                reset_results["azure_search"]["message"] = f"Error: {str(search_ex)}"
                logger.error(
                    f"[ERROR] [FACTORY RESET] Azure Search deletion error: {search_ex}"
                )

            file_processor.blob_service.update_default_config()
            # Reload config in memory
            config.reload_config()

            logger.warning("[WARNING] [FACTORY RESET] Factory reset completed!")
            logger.warning(f"[WARNING] [FACTORY RESET] Results: {reset_results}")

            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "message": "Factory reset completed",
                    "results": reset_results,
                    "timestamp": timestamp,
                    "performed_by": current_user.get("sub", "unknown"),
                },
            )
        else:
            logger.error(
                "[ERROR] [FACTORY RESET] Factory reset is disabled (FACTORY_RESET_BOT != true)"
            )
            return JSONResponse(
                status_code=status.HTTP_404_NOT_FOUND,
                content=ErrorResponse(
                    status=404,
                    reason="Factory Reset Disabled",
                    location=location,
                    message="Factory reset is disabled.",
                    timestamp=timestamp,
                ).model_dump(),
            )

    except Exception as ex:
        logger.error(f"[ERROR] [FACTORY RESET] Exception during factory reset: {ex}")
        logger.error(f"[ERROR] [FACTORY RESET] Exception type: {type(ex).__name__}")
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=location,
                message=f"Factory reset failed: {str(ex)}",
                timestamp=timestamp,
            ).model_dump(),
        )


@app.get(
    "/v1/metadata-template",
    responses={
        200: {
            "description": "Excel template file with locked headers",
            "content": {
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": {}
            },
        },
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Download metadata Excel template with locked headers",
    tags=["Configuration"],
)
async def get_metadata_template(request: Request):
    """Generate and download an Excel template with required headers (header row is locked)"""
    try:
        cfg = config  # Use global config instance

        # Get required headers from config with proper validation
        headers = (
            cfg.required_headers
            if hasattr(cfg, "required_headers") and cfg.required_headers
            else ["file_name"]
        )

        # Ensure headers is a list
        if not isinstance(headers, list):
            logger.error(
                f"Invalid required_headers type: {type(headers)}, expected list"
            )
            headers = ["file_name"]

        # Filter out None values and ensure all are strings
        headers = [str(h) for h in headers if h is not None]

        if not headers:
            logger.warning("No valid headers found in config, using default")
            headers = ["file_name"]

        # Create workbook and worksheet
        wb = Workbook()
        ws = wb.active
        ws.title = "Metadata"

        # Write header row with styling
        for col_idx, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            # Style the header
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(
                start_color="4472C4", end_color="4472C4", fill_type="solid"
            )
            # Lock the cell
            cell.protection = Protection(locked=True)

        # Add sample row with unlocked cells
        sample_row = []
        for header in headers:
            if header.lower() == "file_name":
                sample_row.append("example_document.pdf")
            elif "date" in header.lower():
                sample_row.append("2025-01-01")
            elif "level" in header.lower() or "category" in header.lower():
                sample_row.append("sample_value")
            else:
                sample_row.append("")

        for col_idx, value in enumerate(sample_row, start=1):
            cell = ws.cell(row=2, column=col_idx, value=value)
            # Unlock data cells
            cell.protection = Protection(locked=False)

        # Protect the sheet (this locks the header row but allows editing other cells)
        # Don't set password to None - just enable sheet protection without a password
        ws.protection.sheet = True

        # Set protection options
        if hasattr(ws.protection, "formatCells"):
            ws.protection.formatCells = False
            ws.protection.formatColumns = False
            ws.protection.formatRows = False
            ws.protection.insertColumns = False
            ws.protection.insertRows = False
            ws.protection.deleteColumns = False
            ws.protection.deleteRows = False

        # Adjust column widths
        for col_idx in range(1, len(headers) + 1):
            col_letter = ws.cell(row=1, column=col_idx).column_letter
            ws.column_dimensions[col_letter].width = 20

        # Save to bytes
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        # Return as streaming response
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": "attachment; filename=metadata_template.xlsx",
            },
        )

    except Exception as e:
        logger.error(f"Error generating metadata template: {e}")
        logger.exception("Error generating metadata template - full traceback")
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=str(request.url),
                message=f"Failed to generate metadata template: {str(e)}",
                timestamp=datetime.now(timezone.utc).isoformat(),
            ).model_dump(),
        )


@app.post(
    "/v1/chat/history",
    response_model=ChatHistoryApiResponse,
    summary="Get chat history",
    tags=["Chat History"],
)
async def get_chat_history(request: ChatHistoryQuery, http_request: Request):
    # Get current user from middleware
    current_user = get_current_user_from_request(http_request)
    """Get chat history from the external service"""
    ()
    try:
        # Verify user access if sessionID or userID provided
        if (
            current_user
            and request.UserID
            and current_user.get("UserID") != request.UserID
        ):
            raise HTTPException(status_code=403, detail="Access denied")

        result = chat_history_service.get_user_history(request)

        return ChatHistoryApiResponse(
            success=result["success"],
            message=(
                "Chat history retrieved"
                if result["success"]
                else f"Retrieval failed: {result.get('error')}"
            ),
            data=result.get("data"),
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting chat history: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Failed to get chat history: {str(e)}"
        )


@app.post(
    "/v1/chat/feedback",
    response_model=ChatHistoryApiResponse,
    summary="Update message feedback",
    tags=["Chat History"],
)
async def update_message_feedback(
    request: FeedbackUpdateRequest, http_request: Request
):
    # Get current user from middleware
    current_user = get_current_user_from_request(http_request)
    """Update feedback for a chat message"""
    ()
    try:
        # Get user_id from current_user (using lowercase 'user_id' as set by middleware)
        current_user_id = (
            current_user.get("user_id")
            or current_user.get("sub")
            or current_user.get("oid")
        )

        # If UserID is not provided in request, use the one from auth token
        if not request.UserID:
            request.UserID = current_user_id

        # Override bot_id with config value
        request.BotID = config.bot_id

        result = chat_history_service.update_feedback(request)

        response = ChatHistoryApiResponse(
            success=result["success"],
            message=(
                "Feedback updated"
                if result["success"]
                else f"Update failed: {result.get('error')}"
            ),
            data=result.get("data"),
        )

        return response

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating feedback: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Failed to update feedback: {str(e)}"
        )


@app.get(
    "/v1/bots/{bot_id}/statistics",
    responses={
        200: {"description": "Bot statistics retrieved successfully"},
        400: {"model": ErrorResponse, "description": "Bad Request"},
        401: {"model": ErrorResponse, "description": "Unauthorized"},
        403: {
            "model": ErrorResponse,
            "description": "Forbidden - Authentication required",
        },
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Get bot statistics for a specific time period (Authentication required)",
    tags=["Bot Statistics"],
)
async def get_bot_statistics(request: Request, bot_id: str, time_range: str = "today"):
    """
    Get bot statistics for a specific time period.

    Args:
        bot_id: Bot identifier
        time_range: Time period filter - 'today', 'this_week', or 'this_month' (default: 'today')

    Returns:
        JSON response with bot statistics including total messages, active users, sessions, etc.
    """

    location = str(request.url)
    timestamp = datetime.now(timezone.utc).isoformat()

    try:
        if not chat_history_service:
            logger.error("[ERROR] [GET BOT STATISTICS] Chat history service not available")
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=location,
                    message="Chat history service not available",
                    timestamp=timestamp,
                ).model_dump(),
            )

        # Call the chat history service
        result = chat_history_service.get_bot_statistics(bot_id, time_range)

        if result["success"]:
            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "success": True,
                    "data": result["data"],
                    "timestamp": timestamp,
                },
            )
        else:
            logger.error(
                f"[ERROR] [GET BOT STATISTICS] Failed to retrieve statistics: {result.get('error')}"
            )
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=location,
                    message=result.get("error", "Failed to retrieve bot statistics"),
                    timestamp=timestamp,
                ).model_dump(),
            )

    except Exception as ex:
        logger.error(f"[ERROR] [GET BOT STATISTICS] Exception: {ex}")
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=location,
                message=str(ex),
                timestamp=timestamp,
            ).model_dump(),
        )


@app.post(
    "/v1/chat/export",
    responses={
        200: {"description": "Chat history exported successfully"},
        400: {"model": ErrorResponse, "description": "Bad Request"},
        401: {"model": ErrorResponse, "description": "Unauthorized"},
        403: {
            "model": ErrorResponse,
            "description": "Forbidden - Authentication required",
        },
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
    },
    summary="Export chat history",
    tags=["Chat History"],
)
async def export_chat_history(request: Request):
    # Get current user from middleware
    current_user = get_current_user_from_request(request)
    location = str(request.url)
    timestamp = datetime.now(timezone.utc).isoformat()
    bot_id = config.bot_id
    user_id = current_user.get("user_id") or current_user.get("userID")
    body = (
        await request.json()
        if request.headers.get("content-type") == "application/json"
        else {}
    )
    period = body.get("period", "all")
    export_format = body.get("format", "json")
    """Export chat history for download"""
    try:
        if not chat_history_service:
            logger.error("[ERROR] [Export Chat history] Chat history service not available")
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=location,
                    message="Chat history service not available",
                    timestamp=timestamp,
                ).model_dump(),
            )

        query = ChatExportRequest(
            BotID=bot_id,
            UserID=user_id,
            period=period,
        )
        result = chat_history_service.get_user_history_export(query)
        if result["success"]:
            logger.info(
                f"[INFO] [CHAT HISTORY EXPORT ENDPOINT BE] Retrieved {len(result['data'])} messages for export"
            )
        else:
            logger.error(
                f"[ERROR] [CHAT HISTORY EXPORT ENDPOINT BE] Failed to retrieve chat history: {result.get('error')}"
            )
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=location,
                    message=f"Failed to retrieve chat history: {result.get('error')}",
                    timestamp=timestamp,
                ).model_dump(),
            )

        chat_data = result["data"]

        # Extract items from the result (now grouped by SessionID)
        items = chat_data.get("items", [])
        total_count = chat_data.get("total_count", 0)

        # Handle case where there's no chat history
        if total_count == 0 or not items:
            logger.info(
                f"[INFO] [CHAT HISTORY EXPORT] No chat history found for user {user_id} (period: {period})"
            )

            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "items": [],
                    "total_count": 0,
                    "period": period,
                    "message": "No chat history available for the specified period",
                },
            )

        if export_format.lower() == "csv":
            # Convert to CSV format
            output = io.StringIO()
            fieldnames = [
                "session_number",
                "query",
                "response",
                "citations",
            ]
            writer = csv.DictWriter(output, fieldnames=fieldnames)
            writer.writeheader()
            session_number = 1
            # Flatten the grouped items for CSV export
            for session_group in items[::-1]:
                for item in session_group[::-1]:
                    # Format citations as a readable string
                    citations = item.get("citations", [])
                    citations_str = ""
                    if citations:
                        citation_parts = []
                        for idx, citation in enumerate(citations, 1):
                            file_name = citation.get("file_name", "Unknown")
                            page = citation.get("page_number", "N/A")
                            citation_parts.append(f"[{idx}] {file_name} (Page {page})")
                        citations_str = "; ".join(citation_parts)

                    writer.writerow(
                        {
                            "session_number": session_number,
                            "query": item.get("query", ""),
                            "response": item.get("response", ""),
                            "citations": citations_str,
                        }
                    )
                session_number += 1
            # Return CSV as a streaming response
            return StreamingResponse(
                io.BytesIO(output.getvalue().encode("utf-8")),
                media_type="text/csv",
                headers={
                    "Content-Disposition": f"attachment; filename=chat_history_{user_id}_{period}.csv"
                },
            )

        elif export_format.lower() == "word" or export_format.lower() == "docx":
            # Convert to Word document format
            doc = Document()

            # Add title
            title = doc.add_heading("Chat History Export", 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add chat sessions
            session_number = 1
            for session_group in items[::-1]:
                if session_group:
                    # Session header
                    session_heading = doc.add_heading(f"Session {session_number}", 1)
                    session_heading.style.font.color.rgb = RGBColor(0, 0, 139)

                    # Add messages in the session
                    for item in session_group[::-1]:
                        # Query
                        query_para = doc.add_paragraph()
                        query_run = query_para.add_run(f"Q: {item.get('query', '')}")
                        query_run.bold = True
                        query_run.font.color.rgb = RGBColor(0, 100, 0)

                        # Response
                        response_para = doc.add_paragraph()
                        response_run = response_para.add_run(
                            f"A: {item.get('response', '')}"
                        )
                        response_run.font.color.rgb = RGBColor(0, 0, 0)

                        # Citations (if any)
                        citations = item.get("citations", [])
                        if citations:
                            citations_para = doc.add_paragraph()
                            citations_run = citations_para.add_run("Citations:")
                            citations_run.bold = True
                            citations_run.font.size = Pt(9)
                            citations_run.font.color.rgb = RGBColor(70, 70, 70)

                            for idx, citation in enumerate(citations, 1):
                                file_name = citation.get("file_name", "Unknown")
                                page = citation.get("page_number", "N/A")
                                citation_para = doc.add_paragraph(
                                    f"  [{idx}] {file_name} (Page {page})",
                                    style="List Bullet",
                                )
                                citation_para.paragraph_format.left_indent = Pt(30)
                                citation_run = citation_para.runs[0]
                                citation_run.font.size = Pt(9)
                                citation_run.font.color.rgb = RGBColor(100, 100, 100)

                        doc.add_paragraph()  # Empty line between messages

                    session_number += 1
                    doc.add_page_break()  # New page for each session

            # Save to BytesIO
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            return StreamingResponse(
                output,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename=chat_history_{user_id}_{period}.docx"
                },
            )

        elif export_format.lower() == "pdf":
            # Convert to PDF format
            output = io.BytesIO()
            doc = SimpleDocTemplate(output, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()

            # Create custom styles
            title_style = ParagraphStyle(
                "CustomTitle",
                parent=styles["Heading1"],
                fontSize=24,
                textColor="darkblue",
                alignment=TA_CENTER,
                spaceAfter=30,
            )

            session_style = ParagraphStyle(
                "SessionHeading",
                parent=styles["Heading2"],
                fontSize=16,
                textColor="blue",
                spaceAfter=12,
            )

            query_style = ParagraphStyle(
                "Query",
                parent=styles["Normal"],
                fontSize=11,
                textColor="green",
                leftIndent=20,
                spaceAfter=6,
            )

            response_style = ParagraphStyle(
                "Response",
                parent=styles["Normal"],
                fontSize=10,
                leftIndent=20,
                spaceAfter=6,
            )

            meta_style = ParagraphStyle(
                "Meta",
                parent=styles["Normal"],
                fontSize=8,
                textColor="gray",
                leftIndent=20,
                spaceAfter=12,
            )

            # Add title
            story.append(Paragraph("Chat History Export", title_style))
            story.append(Spacer(1, 0.2 * inch))

            # Add chat sessions
            session_number = 1
            for session_group in items[::-1]:
                if session_group:
                    # Session header
                    story.append(Paragraph(f"Session {session_number}", session_style))
                    story.append(Spacer(1, 0.1 * inch))

                    # Add messages in the session
                    for item in session_group[::-1]:
                        # Query
                        query_text = (
                            item.get("query", "")
                            .replace("<", "&lt;")
                            .replace(">", "&gt;")
                        )
                        story.append(Paragraph(f"<b>Q:</b> {query_text}", query_style))

                        # Response
                        response_text = (
                            item.get("response", "")
                            .replace("<", "&lt;")
                            .replace(">", "&gt;")
                        )
                        story.append(
                            Paragraph(f"<b>A:</b> {response_text}", response_style)
                        )

                        # Citations (if any)
                        citations = item.get("citations", [])
                        if citations:
                            story.append(Paragraph("<b>Citations:</b>", meta_style))
                            for idx, citation in enumerate(citations, 1):
                                file_name = citation.get("file_name", "Unknown")
                                page = citation.get("page_number", "N/A")
                                citation_text = f"[{idx}] {file_name} (Page {page})"
                                story.append(
                                    Paragraph(
                                        f"&nbsp;&nbsp;{citation_text}", meta_style
                                    )
                                )

                        story.append(Spacer(1, 0.15 * inch))

                    session_number += 1
                    story.append(PageBreak())  # New page for each session

            # Build PDF
            doc.build(story)
            output.seek(0)
            return StreamingResponse(
                output,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename=chat_history_{user_id}_{period}.pdf"
                },
            )

        elif export_format.lower() == "json":
            # Return JSON with the grouped structure
            return JSONResponse(
                status_code=status.HTTP_200_OK,
                content={
                    "items": items,  # Grouped by SessionID
                    "total_count": total_count,
                    "period": period,
                },
            )
        else:
            return JSONResponse(
                status_code=status.HTTP_400_BAD_REQUEST,
                content=ErrorResponse(
                    status=400,
                    reason="Bad Request",
                    location=location,
                    message=f"Unsupported format: {export_format}. Supported formats: 'json', 'csv', 'word', 'docx', 'pdf'.",
                    timestamp=timestamp,
                ).model_dump(),
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in exporting chat history: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to export chat: {str(e)}")


@app.get(
    "/v1/session/{session_id}",
    response_model=ChatHistoryApiResponse,
    summary="Get session details",
    tags=["Chat History"],
)
async def get_session_details(session_id: str, request: Request):
    """
    Get details for a private session. Requires authentication.
    """
    try:
        # Use chat history service to get session details
        if not chat_history_service:
            raise HTTPException(
                status_code=503, detail="Chat history service not available"
            )

        current_user = get_current_user_from_request(request)
        user_id = current_user.get("user_id") or current_user.get("userID")
        if not user_id:
            raise HTTPException(status_code=401, detail="User ID not found in token")

        result = chat_history_service.get_user_session(
            userID=user_id, sessionID=session_id, bot_id=config.bot_id
        )
        if not result["success"]:
            logger.warning(
                "%s session_details.fetch_failed session_id=%s error=%s",
                BACKEND_EXCEPTION_TAG,
                session_id,
                result.get("error"),
            )
            raise HTTPException(
                status_code=500, detail=result.get("error", "Failed to get session")
            )

        all_messages = result["data"].get("messages", [])
        if not all_messages:
            raise HTTPException(status_code=404, detail="Session not found")

        first_message = all_messages[0]
        session_user_id = first_message["UserID"]

        if user_id != session_user_id:
            raise HTTPException(
                status_code=403,
                detail="Access denied: Session belongs to different user",
            )

        session_data = {
            "SessionID": session_id,
            "UserID": session_user_id,
            "bot_id": config.bot_id,
            "created_at": (
                min(msg["created_at"] for msg in all_messages) if all_messages else None
            ),
            "last_activity": (
                max(msg["created_at"] for msg in all_messages) if all_messages else None
            ),
            "message_count": len(all_messages),
            "messages": all_messages,
            "metadata": {
                "total_messages": len(all_messages),
                "first_query": (all_messages[-1]["query"] if all_messages else None),
                "latest_query": (all_messages[0]["query"] if all_messages else None),
            },
            "is_shared": False,
        }

        return ChatHistoryApiResponse(
            success=True, message="Session details retrieved", data=session_data
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting session details: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get session: {str(e)}")


@app.post(
    "/v1/session/{session_id}/share",
    response_model=SessionShareResponse,
    summary="Create shareable link for a session",
    tags=["Session Sharing"],
)
async def create_session_share(
    session_id: str,
    request: Request,
    share_request: Optional[SessionShareRequest] = None,
):
    """
    Create a shareable link for a session.
    This marks the session as PUBLIC (shareable).
    Requires authentication - only session owner can share.
    """
    # Get current user from middleware
    current_user = get_current_user_from_request(request)
    user_id = current_user.get("user_id") or current_user.get("userID")
    if not user_id:
        raise HTTPException(status_code=401, detail="User ID not found in token")

    try:
        # Verify user owns the session
        if not chat_history_service:
            raise HTTPException(
                status_code=503, detail="Chat history service not available"
            )

        # Check if session exists and user owns it
        result = chat_history_service.get_user_session(
            userID=user_id, sessionID=session_id, bot_id=config.bot_id
        )

        if not result["success"]:
            logger.warning(
                "%s session_share.create.fetch_failed session_id=%s error=%s",
                BACKEND_EXCEPTION_TAG,
                session_id,
                result.get("error"),
            )
            raise HTTPException(
                status_code=404, detail="Session not found or access denied"
            )

        # Get expiration days from request or use default
        expires_in_days = share_request.expires_in_days if share_request else 30

        # Create share token
        share_info = session_share_service.create_share_token(
            session_id=session_id,
            user_id=user_id,
            bot_id=config.bot_id,
            expires_in_days=expires_in_days,
        )

        return SessionShareResponse(
            success=True,
            share_token=share_info["share_token"],
            expires_at=share_info["expires_at"],
            created_at=share_info["created_at"],
            message="Session share link created successfully",
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating session share: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Failed to create share link: {str(e)}"
        )


@app.delete(
    "/v1/session/{session_id}/share",
    summary="Revoke share token (make session private)",
    tags=["Session Sharing"],
)
async def revoke_session_share(session_id: str, request: Request):
    """
    Revoke share token for a session.
    This marks the session as PRIVATE again.
    Only session owner can revoke.
    """
    # Get current user from middleware
    current_user = get_current_user_from_request(request)
    user_id = current_user.get("user_id") or current_user.get("userID")
    if not user_id:
        raise HTTPException(status_code=401, detail="User ID not found in token")

    try:
        # Verify user owns the session
        if not chat_history_service:
            raise HTTPException(
                status_code=503, detail="Chat history service not available"
            )

        # Check if session exists and user owns it
        result = chat_history_service.get_user_session(
            userID=user_id, sessionID=session_id, bot_id=config.bot_id
        )

        if not result["success"]:
            logger.warning(
                "%s session_share.revoke.fetch_failed session_id=%s error=%s",
                BACKEND_EXCEPTION_TAG,
                session_id,
                result.get("error"),
            )
            raise HTTPException(
                status_code=404, detail="Session not found or access denied"
            )

        # Revoke share token
        revoked = session_share_service.revoke_share_token(
            session_id=session_id, user_id=user_id
        )

        if not revoked:
            return ChatHistoryApiResponse(
                success=False,
                message="No active share token found for this session",
                data=None,
            )

        return ChatHistoryApiResponse(
            success=True,
            message="Session share link revoked successfully. Session is now private.",
            data=None,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error revoking session share: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Failed to revoke share link: {str(e)}"
        )


@app.get(
    "/v1/public_session/{session_id}",
    response_model=ChatHistoryApiResponse,
    summary="Get public session details",
    tags=["Session Sharing"],
)
async def get_public_session(session_id: str, request: Request):
    """
    Get details for a public/shared session. No authentication required.
    """
    try:
        if not chat_history_service:
            raise HTTPException(
                status_code=503, detail="Chat history service not available"
            )

        result = chat_history_service.get_public_session(
            session_id=session_id, bot_id=config.bot_id
        )

        if not result.get("success", False):
            logger.warning(
                "%s session_share.public.fetch_failed session_id=%s error=%s",
                BACKEND_EXCEPTION_TAG,
                session_id,
                result.get("error"),
            )
            raise HTTPException(
                status_code=404, detail="Session not found or not public"
            )

        all_messages = result["data"].get("items", [])
        if not all_messages:
            raise HTTPException(status_code=404, detail="Session not found")

        first_message = all_messages[0]
        session_user_id = first_message["UserID"]

        session_data = {
            "SessionID": session_id,
            "UserID": session_user_id,
            "bot_id": config.bot_id,
            "created_at": (
                min(msg["created_at"] for msg in all_messages) if all_messages else None
            ),
            "last_activity": (
                max(msg["created_at"] for msg in all_messages) if all_messages else None
            ),
            "message_count": len(all_messages),
            "messages": all_messages,
            "metadata": {
                "total_messages": len(all_messages),
                "first_query": (all_messages[-1]["query"] if all_messages else None),
                "latest_query": (all_messages[0]["query"] if all_messages else None),
            },
            "is_shared": True,
        }

        return ChatHistoryApiResponse(
            success=True, message="Public session details retrieved", data=session_data
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting public session: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Failed to get public session: {str(e)}"
        )


@app.get(
    "/v1/session/{session_id}/share",
    response_model=SessionShareInfoResponse,
    summary="Get share info for a session",
    tags=["Session Sharing"],
)
async def get_session_share_info(session_id: str, request: Request):
    """
    Check if a session is currently shared and get share URL.
    Only session owner can check.
    """
    # Get current user from middleware
    current_user = get_current_user_from_request(request)
    user_id = current_user.get("user_id") or current_user.get("userID")
    if not user_id:
        raise HTTPException(status_code=401, detail="User ID not found in token")

    try:
        # Verify user owns the session
        if not chat_history_service:
            raise HTTPException(
                status_code=503, detail="Chat history service not available"
            )

        # Check if session exists and user owns it
        result = chat_history_service.get_user_session(
            userID=user_id, sessionID=session_id, bot_id=config.bot_id
        )

        if not result["success"]:
            logger.warning(
                "%s session_share.info.fetch_failed session_id=%s error=%s",
                BACKEND_EXCEPTION_TAG,
                session_id,
                result.get("error"),
            )
            raise HTTPException(
                status_code=404, detail="Session not found or access denied"
            )

        # Get share info
        share_info = session_share_service.get_session_share_info(
            session_id=session_id, user_id=user_id
        )

        if share_info:
            return SessionShareInfoResponse(
                success=True,
                is_shared=True,
                share_token=share_info["share_token"],
                expires_at=share_info["expires_at"],
                created_at=share_info["created_at"],
                message="Session is currently shared",
            )
        else:
            return SessionShareInfoResponse(
                success=True,
                is_shared=False,
                share_token=None,
                expires_at=None,
                created_at=None,
                message="Session is private (not shared)",
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting session share info: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Failed to get share info: {str(e)}"
        )


@app.websocket("/v1/ws/status")
async def websocket_status_endpoint(websocket: WebSocket):
    """
    WebSocket endpoint for real-time status updates

    Clients can send subscription messages in the format:
    {"action": "subscribe", "work_id": "work_id_here"}

    And will receive status updates in the format:
    {"type": "status_update", "work_id": "work_id", "data": {...status data...}}
    """
    await connection_manager.connect(websocket)

    try:
        while True:
            # Wait for messages from client
            data = await websocket.receive_text()
            try:
                message = json.loads(data)
                action = message.get("action")

                if action == "subscribe":
                    work_id = message.get("work_id")
                    if work_id:
                        connection_manager.subscribe_to_work_id(websocket, work_id)

                        # Send current status immediately upon subscription
                        upload_record = file_processor.get_upload_info(work_id)
                        if upload_record:
                            await connection_manager.send_personal_message(
                                {
                                    "type": "status_update",
                                    "work_id": work_id,
                                    "data": upload_record,
                                },
                                websocket,
                            )
                        else:
                            await connection_manager.send_personal_message(
                                {
                                    "type": "error",
                                    "message": f"Work ID {work_id} not found",
                                },
                                websocket,
                            )
                    else:
                        await connection_manager.send_personal_message(
                            {
                                "type": "error",
                                "message": "work_id is required for subscription",
                            },
                            websocket,
                        )

                elif action == "ping":
                    # Respond to ping with pong
                    await connection_manager.send_personal_message(
                        {
                            "type": "pong",
                            "timestamp": datetime.now(timezone.utc).isoformat(),
                        },
                        websocket,
                    )

                else:
                    await connection_manager.send_personal_message(
                        {"type": "error", "message": f"Unknown action: {action}"},
                        websocket,
                    )

            except json.JSONDecodeError:
                await connection_manager.send_personal_message(
                    {"type": "error", "message": "Invalid JSON format"}, websocket
                )
            except Exception as e:
                logger.error(f"Error processing WebSocket message: {e}")
                await connection_manager.send_personal_message(
                    {"type": "error", "message": f"Error processing message: {str(e)}"},
                    websocket,
                )

    except WebSocketDisconnect:
        connection_manager.disconnect(websocket)
    except Exception as e:
        logger.error(f"WebSocket error: {e}")
        connection_manager.disconnect(websocket)


@app.put(
    "/v1/updateconfig",
    response_model=dict,
    responses={
        200: {"description": "Bot configuration updated successfully"},
        400: {"model": ErrorResponse, "description": "Bad Request"},
        500: {"model": ErrorResponse, "description": "Internal Server Error"},
        401: {"model": ErrorResponse, "description": "Unauthorized"},
        403: {
            "model": ErrorResponse,
            "description": "Forbidden - Admin access required",
        },
    },
    summary="Update bot configuration (Admin only)",
    tags=["Configuration"],
)
async def update_bot_configuration(
    request: Request,
    new_config: dict = Body(..., description="New bot configuration data"),
):
    """
    Admin endpoint to update the bot parameters in config.
    Returns the updated theme configuration for frontend TOML generation.
    """
    # Get current user from middleware (admin access already validated by middleware)
    try:
        current_user = get_current_user_from_request(request)
        user_id = (
            current_user.get("user_id") or current_user.get("sub")
            if current_user
            else None
        )
        if not user_id:
            logger.error("[ERROR] [UPDATE CONFIG] No authenticated user found")
            return JSONResponse(
                status_code=status.HTTP_401_UNAUTHORIZED,
                content=ErrorResponse(
                    status=401,
                    reason="Unauthorized",
                    location=str(request.url),
                    message="Authentication required",
                    timestamp=datetime.now(timezone.utc).isoformat(),
                ).model_dump(),
            )

    except Exception as auth_ex:
        logger.error(f"[ERROR] [UPDATE CONFIG] Authentication error: {auth_ex}")
        return JSONResponse(
            status_code=status.HTTP_401_UNAUTHORIZED,
            content=ErrorResponse(
                status=401,
                reason="Unauthorized",
                location=str(request.url),
                message="Authentication failed",
                timestamp=datetime.now(timezone.utc).isoformat(),
            ).model_dump(),
        )

    try:
        cfg = config  # Use global config instance

        # Validate that storage is configured
        if not cfg.azure_storage_account_name or not cfg.azure_storage_container_name:
            logger.error("[ERROR] [UPDATE CONFIG] Azure Storage not configured")
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=str(request.url),
                    message="Azure Storage not configured",
                    timestamp=datetime.now(timezone.utc).isoformat(),
                ).model_dump(),
            )

        # Get current bot config from blob storage
        try:
            credential = DefaultAzureCredential()
            account_url = (
                f"https://{cfg.azure_storage_account_name}.blob.core.windows.net"
            )
            blob_service_client = BlobServiceClient(
                account_url=account_url, credential=credential
            )
            container_client = blob_service_client.get_container_client(
                cfg.azure_storage_container_name
            )
            config_blob_name = "config.json"
            blob_client = container_client.get_blob_client(config_blob_name)

            # Download current config or create new one
            if blob_client.exists():
                blob_data = blob_client.download_blob().readall()
                current_config = json.loads(blob_data.decode("utf-8"))
            else:
                current_config = {}

            # --- Completely flatten look_and_feel into root and remove the field ---
            if "look_and_feel" in current_config and isinstance(
                current_config["look_and_feel"], dict
            ):
                lf = current_config["look_and_feel"]

                # Move all look_and_feel fields to root (root always gets overridden)
                for k, v in lf.items():
                    current_config[k] = v

                # Remove the look_and_feel field entirely
                del current_config["look_and_feel"]
            # --- Handle base64 images directly from payload ---
            image_inputs = [
                ("logo_image_base64", "logo_filename", "logo_file_name", "logo"),
                (
                    "bot_icon_image_base64",
                    "bot_icon_filename",
                    "bot_icon_file_name",
                    "bot_icon",
                ),
                (
                    "user_icon_image_base64",
                    "user_icon_filename",
                    "user_icon_file_name",
                    "user_icon",
                ),
            ]

            for b64_key, name_key, cfg_key, img_type in image_inputs:
                b64_value = new_config.get(b64_key)
                if not b64_value:
                    continue

                # Decode base64 content
                raw_b64 = b64_value.split(",", 1)[-1] if "," in b64_value else b64_value
                file_bytes = base64.b64decode(raw_b64)

                # Infer extension from filename
                orig_name = new_config.get(name_key) or f"{img_type}.png"
                ext = Path(orig_name).suffix or ".png"
                if not ext.startswith("."):
                    ext = f".{ext}"

                # Generate a unique blob name
                safe_type = re.sub(r"[^a-z0-9_-]", "", img_type.lower())[:40] or "image"
                blob_name = f"{safe_type}{ext}"

                # Infer content-type
                if ext.lower() in [".jpg", ".jpeg"]:
                    content_type = "image/jpeg"
                elif ext.lower() == ".gif":
                    content_type = "image/gif"
                elif ext.lower() == ".webp":
                    content_type = "image/webp"
                else:
                    content_type = "image/png"

                # Upload the file bytes to blob
                if file_processor and getattr(file_processor, "blob_service", None):
                    file_processor.blob_service.upload_bytes(
                        blob_name, file_bytes, content_type=content_type, metadata=None
                    )
                else:
                    container_client.get_blob_client(blob_name).upload_blob(
                        file_bytes,
                        overwrite=True,
                        content_settings=ContentSettings(content_type=content_type),
                    )

                # Update filename
                current_config[cfg_key] = blob_name
                cfg.bot_config[cfg_key] = blob_name

            # Deep-merge for nested dicts; overwrite for scalars
            for k, v in new_config.items():
                if k in ("filters", "filter_mapping") and isinstance(v, dict):
                    base = current_config.get(k) or {}
                    if not isinstance(base, dict):
                        base = {}
                    # Only update provided keys; leave others intact
                    for sk, sv in v.items():
                        base[sk] = sv
                    current_config[k] = base
                else:
                    current_config[k] = v

            # Mirror into global cfg.bot_config as well
            for k, v in new_config.items():
                if k in ("filters", "filter_mapping") and isinstance(v, dict):
                    base = cfg.bot_config.get(k) or {}
                    if not isinstance(base, dict):
                        base = {}
                    for sk, sv in v.items():
                        base[sk] = sv
                    cfg.bot_config[k] = base
                else:
                    cfg.bot_config[k] = v

            allowed_top_keys = {
                "has_filters",
                "filters",
                "required_headers",
                "filter_mapping",
                "bot_name",
                "version",
                "language",
                "about_text",
                "disclaimer_text",
                "primary_color",
                "secondary_background_color",
                "background_color",
                "text_color",
                "font_family",
                "font_size",
                "welcome_message",
                "default_response",
                "logo_file_name",
                "bot_icon_file_name",
                "user_icon_file_name",
                "feedback_contact_name",
                "feedback_contact_email",
                "faq",
                "system_prompt",
                "external_links",
            }
            current_config = {
                k: v for k, v in current_config.items() if k in allowed_top_keys
            }
            cfg.bot_config = {
                k: v for k, v in cfg.bot_config.items() if k in allowed_top_keys
            }

            # Upload updated config.json back to blob
            updated_config_json = json.dumps(current_config, indent=2)
            blob_client.upload_blob(
                updated_config_json.encode("utf-8"),
                overwrite=True,
                content_settings=ContentSettings(content_type="application/json"),
            )

            # Reload the config to ensure system_prompt and other settings are updated
            logger.info("[INFO] [UPDATE CONFIG] Reloading configuration after update")
            cfg.reload_config()
            logger.info(
                f"[INFO] [UPDATE CONFIG] Config reloaded. New system_prompt: {cfg.system_prompt[:100] if cfg.system_prompt else 'None'}..."
            )

            # --- Build response (no helper functions) ---
            # Read image bytes again from blob, encode as base64 string
            images_dict = {}
            for key in ["logo_file_name", "bot_icon_file_name", "user_icon_file_name"]:
                fname = current_config.get(key)
                if not fname:
                    images_dict[f"{key.split('_file_name')[0]}_base64"] = None
                    continue
                try:
                    blob_data = (
                        container_client.get_blob_client(fname)
                        .download_blob()
                        .readall()
                    )
                    images_dict[f"{key.split('_file_name')[0]}_base64"] = (
                        base64.b64encode(blob_data).decode("ascii")
                    )
                except Exception:
                    images_dict[f"{key.split('_file_name')[0]}_base64"] = None

            # Construct a unified config response
            has_filters = current_config.get("has_filters", False)
            response_config = {
                "has_filters": has_filters,
                "filters": current_config.get("filters", {}) if has_filters else {},
                "filter_mapping": (
                    current_config.get("filter_mapping", {}) if has_filters else {}
                ),
                "bot_name": current_config.get("bot_name"),
                "version": current_config.get("version"),
                "language": current_config.get("language"),
                "about_text": current_config.get("about_text"),
                "disclaimer_text": current_config.get("disclaimer_text"),
                "primary_color": current_config.get("primary_color"),
                "secondary_background_color": current_config.get(
                    "secondary_background_color"
                ),
                "background_color": current_config.get("background_color"),
                "text_color": current_config.get("text_color"),
                "font_family": current_config.get("font_family"),
                "font_size": current_config.get("font_size"),
                "welcome_message": current_config.get("welcome_message"),
                "default_response": current_config.get("default_response"),
                "feedback_contact_name": current_config.get("feedback_contact_name"),
                "feedback_contact_email": current_config.get("feedback_contact_email"),
                "faq": current_config.get("faq", []),
                "system_prompt": current_config.get("system_prompt"),
                "external_links": current_config.get("external_links", []),
                "images": images_dict,
            }

            return {
                "status": "success",
                "message": "Bot configuration updated successfully",
                "config": response_config,
                "updated_by": user_id,
                "timestamp": datetime.now(timezone.utc).isoformat(),
            }
        except json.JSONDecodeError as e:
            logger.error(f"[ERROR] [UPDATE CONFIG] Invalid JSON in config.json: {e}")
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=str(request.url),
                    message=f"Invalid JSON in configuration file: {str(e)}",
                    timestamp=datetime.now(timezone.utc).isoformat(),
                ).model_dump(),
            )
        except Exception as e:
            logger.error(f"[ERROR] [UPDATE CONFIG] Error updating config: {e}")
            return JSONResponse(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                content=ErrorResponse(
                    status=500,
                    reason="Internal Server Error",
                    location=str(request.url),
                    message=f"Failed to update configuration: {str(e)}",
                    timestamp=datetime.now(timezone.utc).isoformat(),
                ).model_dump(),
            )

    except Exception as e:
        logger.error(f"[ERROR] [UPDATE CONFIG] Error in update_config endpoint: {e}")
        return JSONResponse(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            content=ErrorResponse(
                status=500,
                reason="Internal Server Error",
                location=str(request.url),
                message=f"Internal server error: {str(e)}",
                timestamp=datetime.now(timezone.utc).isoformat(),
            ).model_dump(),
        )


@app.websocket("/v1/ws/status/{work_id}")
async def websocket_specific_status(websocket: WebSocket, work_id: str):
    """
    WebSocket endpoint for a specific work_id status updates
    Automatically subscribes to the specified work_id
    """
    await connection_manager.connect(websocket)
    connection_manager.subscribe_to_work_id(websocket, work_id)

    try:
        # Send current status immediately
        upload_record = file_processor.get_upload_info(work_id)
        if upload_record:
            await connection_manager.send_personal_message(
                {"type": "status_update", "work_id": work_id, "data": upload_record},
                websocket,
            )
        else:
            await connection_manager.send_personal_message(
                {"type": "error", "message": f"Work ID {work_id} not found"}, websocket
            )

        # Keep connection alive and handle client messages
        while True:
            try:
                data = await websocket.receive_text()
                message = json.loads(data)

                if message.get("action") == "ping":
                    await connection_manager.send_personal_message(
                        {
                            "type": "pong",
                            "timestamp": datetime.now(timezone.utc).isoformat(),
                        },
                        websocket,
                    )

            except json.JSONDecodeError:
                await connection_manager.send_personal_message(
                    {"type": "error", "message": "Invalid JSON format"}, websocket
                )

    except WebSocketDisconnect:
        connection_manager.disconnect(websocket)
    except Exception as e:
        logger.error(f"WebSocket error for work_id {work_id}: {e}")
        connection_manager.disconnect(websocket)


# Cleanup function to properly shutdown background workers
@app.on_event("shutdown")
async def shutdown_event():
    """Cleanup when the server shuts down"""
    logger.info("Shutting down background workers...")
    if file_processor:
        file_processor.stop()
    logger.info("Background workers stopped")
